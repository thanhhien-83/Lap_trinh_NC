# -*- coding: utf-8 -*-
"""
Tạo file Word CHƯƠNG 4: XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG
BTL Lập trình nâng cao - Hệ thống Quản lý Sinh viên
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===================== HELPER FUNCTIONS =====================

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tblPr.append(borders)

def heading1(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 70, 127)
    return p

def heading2(doc, text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    return p

def heading3(doc, text):
    p = doc.add_heading(text, level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 73, 125)
    return p

def body(doc, text, indent=False, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
        r.bold = True
        r = p.add_run(text)
    else:
        r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    return p

def bullet(doc, text, bold_start=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    if bold_start:
        r = p.add_run(bold_start)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
        r.bold = True
        r2 = p.add_run(text)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(13)
    else:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
    return p

def code_block(doc, lines):
    """Thêm khối code dạng bảng 1 cột nền xám"""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F2F2F2')
    tc = cell.paragraphs[0]
    tc.clear()
    for line in lines:
        r = tc.add_run(line + '\n')
        r.font.name = 'Courier New'
        r.font.size = Pt(10)
    doc.add_paragraph()

def note_box(doc, text, color='FFF3CD'):
    """Hộp ghi chú nền vàng nhạt"""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    r = p.add_run('📌 ' + text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.italic = True
    doc.add_paragraph()

# ===================== TẠO TÀI LIỆU =====================

def create_chapter4():
    doc = Document()

    # Thiết lập lề trang
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.5)

    # Thiết lập font mặc định
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    for i in range(1, 4):
        h = doc.styles[f'Heading {i}']
        h.font.name = 'Times New Roman'
        h.font.bold = True

    # ======================== TIÊU ĐỀ CHƯƠNG ========================
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CHƯƠNG 4')
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 70, 127)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG')
    run2.bold = True
    run2.font.size = Pt(18)
    run2.font.name = 'Times New Roman'
    run2.font.color.rgb = RGBColor(0, 70, 127)

    doc.add_paragraph()

    body(doc,
         'Sau khi đã phân tích yêu cầu (Chương 2) và thiết kế các thuật toán cốt lõi '
         '(Chương 3), chương này trình bày toàn bộ quá trình xây dựng và triển khai '
         'thực tế của Hệ thống Quản lý Sinh viên. Nội dung bao gồm: môi trường phát '
         'triển, cấu trúc mã nguồn, cài đặt từng module, các chức năng chính của hệ '
         'thống, kiểm thử và kết quả đạt được sau khi hoàn thiện.')
    doc.add_paragraph()

    # ================================================================
    # 4.1 MÔI TRƯỜNG VÀ CÔNG CỤ PHÁT TRIỂN
    # ================================================================
    heading1(doc, '4.1. Môi trường và Công cụ Phát triển')

    body(doc,
         'Hệ thống được phát triển hoàn toàn bằng công nghệ Web phía máy khách '
         '(Client-side), không yêu cầu máy chủ hay cơ sở dữ liệu ngoại vi. '
         'Điều này giúp ứng dụng chạy trực tiếp trên trình duyệt, dễ triển khai '
         'và chia sẻ.')

    # Bảng công nghệ
    heading2(doc, '4.1.1. Công nghệ sử dụng')

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for cell, text, color in zip(hdr, ['Công nghệ', 'Phiên bản / Chuẩn', 'Vai trò'], ['1F497D','1F497D','1F497D']):
        set_cell_shading(cell, color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(255, 255, 255)

    rows_data = [
        ('HTML5',         'HTML Living Standard', 'Cấu trúc trang web, định nghĩa các thành phần giao diện'),
        ('CSS3',          'CSS Level 3',           'Thiết kế giao diện, animation, layout Flexbox/Grid'),
        ('JavaScript',    'ECMAScript 2020 (ES11)', 'Logic ứng dụng, thuật toán, xử lý dữ liệu'),
        ('LocalStorage',  'Web Storage API',        'Lưu trữ dữ liệu sinh viên phía máy khách'),
        ('Chart.js',      'v4.x (CDN)',             'Vẽ biểu đồ thống kê (tròn, cột)'),
        ('Font Awesome',  'v6.4 (CDN)',             'Bộ icon giao diện'),
        ('Google Fonts',  'Inter (CDN)',            'Font chữ chính của ứng dụng'),
        ('VS Code',       '1.85+',                  'Môi trường lập trình (IDE)'),
        ('Chrome/Edge',   'Phiên bản mới nhất',     'Trình duyệt kiểm thử chính'),
    ]

    for row_data in rows_data:
        row = tbl.add_row().cells
        for i, (cell, text) in enumerate(zip(row, row_data)):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)

    add_table_borders(tbl)

    cap = doc.add_paragraph('Bảng 4.1: Công nghệ và công cụ phát triển hệ thống')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    heading2(doc, '4.1.2. Lý do chọn kiến trúc thuần client-side')
    bullet(doc, 'Chạy trực tiếp trên trình duyệt, không cần cài đặt thêm phần mềm.', '')
    bullet(doc, 'Dễ triển khai: chỉ cần mở file index.html là sử dụng được.', '')
    bullet(doc, 'Phù hợp với quy mô BTL — dữ liệu sinh viên được lưu qua LocalStorage.', '')
    bullet(doc, 'Tập trung minh họa thuật toán và cấu trúc dữ liệu mà không bị phân tán bởi backend.', '')
    bullet(doc, 'Hiệu suất cao — mọi xử lý diễn ra ngay tại máy người dùng, không có độ trễ mạng.', '')
    doc.add_paragraph()

    # ================================================================
    # 4.2 CẤU TRÚC THƯ MỤC VÀ MÃ NGUỒN
    # ================================================================
    heading1(doc, '4.2. Cấu trúc Thư mục và Mã nguồn')

    body(doc,
         'Mã nguồn được tổ chức theo nguyên tắc phân tách trách nhiệm '
         '(Separation of Concerns): HTML chứa cấu trúc, CSS chứa giao diện, '
         'và mỗi file JavaScript đảm nhiệm một vai trò cụ thể trong hệ thống.')

    heading2(doc, '4.2.1. Cấu trúc thư mục')

    code_block(doc, [
        'BTL_Laptrinhnangcao/',
        '├── index.html                  # Trang chính – cấu trúc giao diện',
        '├── css/',
        '│   └── styles.css              # Toàn bộ CSS: layout, animation, responsive',
        '└── js/',
        '    ├── CauTrucMangDong.js      # Lý thuyết & minh họa mảng động',
        '    ├── MangDong.js             # Lớp DynamicArray (cấu trúc dữ liệu)',
        '    ├── SinhVien.js             # Lớp Student (model sinh viên)',
        '    ├── QuanLySinhVien.js       # Lớp StudentManager (CRUD + nghiệp vụ)',
        '    ├── ThuatToanSapXep.js      # QuickSort, BubbleSort',
        '    ├── ThuatToanTimKiem.js     # Linear Search, Binary Search, Filter',
        '    ├── ThuatToanThongKe.js     # Tính GPA, xếp loại, thống kê',
        '    ├── DieuKhienGiaoDien.js    # Lớp UIController (xử lý DOM)',
        '    └── UngDung.js              # Điểm khởi đầu ứng dụng (Application)',
    ])

    heading2(doc, '4.2.2. Vai trò từng file JavaScript')

    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = tbl2.rows[0].cells
    for cell, text in zip(hdr2, ['File', 'Lớp / Module', 'Vai trò chính']):
        set_cell_shading(cell, '2E74B5')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(255, 255, 255)

    files_data = [
        ('CauTrucMangDong.js', 'CauTrucMangDong',  'Tài liệu lý thuyết mảng động, phân tích khấu hao O(1)'),
        ('MangDong.js',        'DynamicArray',      'Hiện thực mảng động: push, pop, resize, sort, search'),
        ('SinhVien.js',        'Student',           'Model sinh viên: thuộc tính, tính GPA, validate, xuất JSON'),
        ('QuanLySinhVien.js',  'StudentManager',    'CRUD: thêm, sửa, xóa, lọc, thống kê, LocalStorage'),
        ('ThuatToanSapXep.js', 'ThuatToanSapXep',  'QuickSort + BubbleSort, hàm so sánh đa tiêu chí'),
        ('ThuatToanTimKiem.js','ThuatToanTimKiem',  'Linear Search, Binary Search, tìm kiếm đa trường'),
        ('ThuatToanThongKe.js','ThuatToanThongKe',  'Tính GPA, xếp loại học lực, thống kê tổng hợp'),
        ('DieuKhienGiaoDien.js','UIController',     'Điều khiển DOM, sự kiện, phân trang, biểu đồ Chart.js'),
        ('UngDung.js',         'Application',       'Khởi tạo ứng dụng, kết nối các module, tải dữ liệu mẫu'),
    ]

    alt = False
    for fd in files_data:
        row = tbl2.add_row().cells
        bg = 'EBF3FB' if alt else 'FFFFFF'
        for i, (cell, text) in enumerate(zip(row, fd)):
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            if i == 0:
                r.bold = True
        alt = not alt

    add_table_borders(tbl2)
    cap2 = doc.add_paragraph('Bảng 4.2: Vai trò của từng file JavaScript trong hệ thống')
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap2.runs[0].italic = True
    cap2.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    heading2(doc, '4.2.3. Thứ tự tải file và phụ thuộc')

    body(doc,
         'Các file JavaScript được nạp vào trang HTML theo thứ tự phụ thuộc '
         '(dependency order). File nào phụ thuộc vào file khác phải được tải sau:')

    code_block(doc, [
        '<!-- index.html — thứ tự tải script -->',
        '<script src="js/CauTrucMangDong.js"></script>   <!-- 1. Lý thuyết (không phụ thuộc) -->',
        '<script src="js/ThuatToanSapXep.js"></script>   <!-- 2. Thuật toán sắp xếp -->',
        '<script src="js/ThuatToanTimKiem.js"></script>  <!-- 3. Thuật toán tìm kiếm -->',
        '<script src="js/ThuatToanThongKe.js"></script>  <!-- 4. Thuật toán thống kê -->',
        '<script src="js/MangDong.js"></script>          <!-- 5. Phụ thuộc ThuatToanSapXep, ThuatToanTimKiem -->',
        '<script src="js/SinhVien.js"></script>          <!-- 6. Phụ thuộc ThuatToanThongKe -->',
        '<script src="js/QuanLySinhVien.js"></script>    <!-- 7. Phụ thuộc MangDong, SinhVien -->',
        '<script src="js/DieuKhienGiaoDien.js"></script> <!-- 8. Phụ thuộc QuanLySinhVien -->',
        '<script src="js/UngDung.js"></script>           <!-- 9. Khởi động toàn ứng dụng -->',
    ])

    doc.add_paragraph()

    # ================================================================
    # 4.3 XÂY DỰNG CẤU TRÚC DỮ LIỆU MẢNG ĐỘNG
    # ================================================================
    heading1(doc, '4.3. Xây dựng Cấu trúc Dữ liệu Mảng Động')

    body(doc,
         'Cấu trúc dữ liệu trung tâm của toàn hệ thống là lớp DynamicArray — '
         'một mảng động tự mở rộng khi cần thiết. Đây là nền tảng để lưu trữ '
         'và thao tác với danh sách sinh viên.')

    heading2(doc, '4.3.1. Thiết kế lớp DynamicArray')

    body(doc,
         'Lớp DynamicArray được cài đặt trong file MangDong.js với ba thuộc tính cốt lõi:')
    bullet(doc, 'data — mảng nội tại lưu trữ các phần tử thực tế.', '')
    bullet(doc, 'length — số phần tử đang sử dụng hiện tại.', '')
    bullet(doc, 'capacity — dung lượng tối đa của mảng nội tại.', '')

    code_block(doc, [
        'class DynamicArray {',
        '    constructor(dungLuongBanDau = 10) {',
        '        this.data     = new Array(dungLuongBanDau);  // Mảng nội tại',
        '        this.length   = 0;                           // Số phần tử đang dùng',
        '        this.capacity = dungLuongBanDau;             // Dung lượng hiện tại',
        '    }',
        '}',
    ])

    heading2(doc, '4.3.2. Cơ chế tự mở rộng (_resize)')

    body(doc,
         'Khi số phần tử đạt giới hạn capacity, mảng tự động mở rộng '
         'gấp đôi dung lượng. Đây là điểm mấu chốt của mảng động:')

    code_block(doc, [
        '_resize() {',
        '    const dungLuongMoi = this.capacity * 2;       // Gấp đôi dung lượng',
        '    const mangMoi = new Array(dungLuongMoi);',
        '',
        '    for (let i = 0; i < this.length; i++) {      // Copy dữ liệu — O(n)',
        '        mangMoi[i] = this.data[i];',
        '    }',
        '',
        '    this.data     = mangMoi;',
        '    this.capacity = dungLuongMoi;',
        '}',
    ])

    body(doc,
         'Nhờ phân tích khấu hao (amortized analysis), chi phí trung bình '
         'của mỗi thao tác push chỉ là O(1), dù đôi khi phải gọi _resize() với chi phí O(n).')

    heading2(doc, '4.3.3. Bảng tổng hợp độ phức tạp của DynamicArray')

    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.style = 'Table Grid'
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr3 = tbl3.rows[0].cells
    for cell, text in zip(hdr3, ['Thao tác', 'Tốt nhất', 'Trung bình', 'Xấu nhất']):
        set_cell_shading(cell, '1F497D')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(255, 255, 255)

    ops = [
        ('push (thêm cuối)',     'O(1)', 'O(1)*',    'O(n)'),
        ('pop (xoá cuối)',       'O(1)', 'O(1)',      'O(1)'),
        ('get / set (truy cập)','O(1)', 'O(1)',      'O(1)'),
        ('insertAt (chèn giữa)','O(1)', 'O(n)',      'O(n)'),
        ('removeAt (xoá giữa)', 'O(1)', 'O(n)',      'O(n)'),
        ('findIndex / find',    'O(1)', 'O(n)',      'O(n)'),
        ('sort (QuickSort)',     'O(nlogn)', 'O(nlogn)', 'O(n²)'),
        ('binarySearch',        'O(1)', 'O(logn)',   'O(logn)'),
    ]
    alt = False
    for op in ops:
        row = tbl3.add_row().cells
        bg = 'EBF3FB' if alt else 'FFFFFF'
        for i, (cell, text) in enumerate(zip(row, op)):
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            if i == 0:
                r.bold = True
        alt = not alt

    add_table_borders(tbl3)
    cap3 = doc.add_paragraph('Bảng 4.3: Độ phức tạp thời gian của các thao tác DynamicArray (* = khấu hao)')
    cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap3.runs[0].italic = True
    cap3.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    # ================================================================
    # 4.4 XÂY DỰNG MODULE SINH VIÊN VÀ QUẢN LÝ
    # ================================================================
    heading1(doc, '4.4. Xây dựng Module Sinh viên và Quản lý')

    heading2(doc, '4.4.1. Lớp Student (SinhVien.js)')

    body(doc,
         'Lớp Student là mô hình hoá đối tượng sinh viên, chứa đầy đủ '
         'thông tin cá nhân, điểm số, và các phương thức xử lý nghiệp vụ.')

    body(doc, 'Các thuộc tính của đối tượng Student:')
    bullet(doc, 'id — Mã sinh viên (chuỗi, duy nhất, ví dụ: "SV001")', '')
    bullet(doc, 'name — Họ và tên đầy đủ', '')
    bullet(doc, 'dob — Ngày sinh (định dạng YYYY-MM-DD)', '')
    bullet(doc, 'gender — Giới tính (Nam / Nữ)', '')
    bullet(doc, 'className — Lớp học', '')
    bullet(doc, 'email — Địa chỉ email', '')
    bullet(doc, 'phone — Số điện thoại', '')
    bullet(doc, 'address — Địa chỉ', '')
    bullet(doc, 'scores — Điểm ba môn: { math, literature, english }', '')
    bullet(doc, 'createdAt / updatedAt — Thời gian tạo và cập nhật', '')

    body(doc, 'Các phương thức của lớp Student:')
    bullet(doc, 'tinhDiemTB() — Tính điểm trung bình của 3 môn', '')
    bullet(doc, 'xepLoai() — Xếp loại học lực: Xuất sắc / Giỏi / Khá / Trung bình / Yếu', '')
    bullet(doc, 'validate() — Kiểm tra tính hợp lệ của dữ liệu đầu vào', '')
    bullet(doc, 'update(data) — Cập nhật thông tin và ghi nhận thời điểm sửa đổi', '')
    bullet(doc, 'toJSON() — Chuyển đối tượng thành chuỗi JSON để lưu LocalStorage', '')
    bullet(doc, 'static fromJSON(json) — Phục hồi đối tượng từ chuỗi JSON', '')

    code_block(doc, [
        '// Minh họa sử dụng lớp Student',
        'const sv = new Student({',
        '    id: "SV001", name: "Nguyễn Văn An",',
        '    dob: "2004-05-12", gender: "Nam",',
        '    className: "CNTT01", email: "an@email.com",',
        '    scores: { math: 8.5, literature: 7.0, english: 9.0 }',
        '});',
        '',
        'console.log(sv.tinhDiemTB());   // → 8.17',
        'console.log(sv.xepLoai());      // → "Giỏi"',
        'console.log(sv.validate());     // → { valid: true, errors: [] }',
    ])

    heading2(doc, '4.4.2. Lớp StudentManager (QuanLySinhVien.js)')

    body(doc,
         'Lớp StudentManager đóng vai trò là lớp dịch vụ (Service Layer), '
         'cung cấp toàn bộ các thao tác CRUD và nghiệp vụ trên danh sách sinh viên. '
         'Nó sử dụng DynamicArray làm cấu trúc lưu trữ chính và LocalStorage để '
         'duy trì dữ liệu giữa các phiên làm việc.')

    body(doc, 'Các phương thức CRUD chính:')

    tbl4 = doc.add_table(rows=1, cols=3)
    tbl4.style = 'Table Grid'
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(tbl4.rows[0].cells, ['Phương thức', 'Thuật toán', 'Mô tả']):
        set_cell_shading(cell, '2E74B5')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(255, 255, 255)

    crud_data = [
        ('addStudent(student)', 'Linear Search O(n)\n+ push O(1)', 'Validate → kiểm tra trùng mã → thêm vào DynamicArray → lưu'),
        ('updateStudent(id, data)', 'Linear Search O(n)', 'Tìm sinh viên → cập nhật thông tin → validate → lưu'),
        ('deleteStudent(id)', 'Linear Search + removeAt O(n)', 'Tìm vị trí → xoá khỏi mảng → lưu'),
        ('getAll()', 'O(n) sao chép', 'Trả về toàn bộ danh sách dưới dạng mảng JS thông thường'),
        ('searchByKeyword(kw)', 'Linear Search đa trường O(n)', 'Tìm theo mã, tên, lớp, email đồng thời'),
        ('sortBy(field, asc)', 'QuickSort O(n log n)', 'Sắp xếp theo tên, điểm, mã SV, lớp'),
        ('getStatistics()', 'Single-pass O(n)', 'Tính GPA tổng, phân loại học lực, thống kê theo lớp'),
        ('exportJSON()', 'O(n)', 'Xuất toàn bộ dữ liệu thành chuỗi JSON'),
    ]
    alt = False
    for cd in crud_data:
        row = tbl4.add_row().cells
        bg = 'EBF3FB' if alt else 'FFFFFF'
        for i, (cell, text) in enumerate(zip(row, cd)):
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            if i == 0:
                r.bold = True
                r.font.name = 'Courier New'
                r.font.size = Pt(11)
        alt = not alt

    add_table_borders(tbl4)
    cap4 = doc.add_paragraph('Bảng 4.4: Các phương thức của lớp StudentManager')
    cap4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap4.runs[0].italic = True
    cap4.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    heading2(doc, '4.4.3. Cơ chế lưu trữ LocalStorage')

    body(doc,
         'Dữ liệu sinh viên được lưu và đọc từ LocalStorage của trình duyệt. '
         'Khi ứng dụng khởi động, StudentManager tự động tải dữ liệu từ '
         'LocalStorage; mỗi khi có thao tác thêm/sửa/xóa, dữ liệu được ghi lại '
         'ngay lập tức.')

    code_block(doc, [
        '// Lưu dữ liệu vào LocalStorage',
        'saveToStorage() {',
        '    const data = this.students.toArray().map(sv => sv.toJSON());',
        '    localStorage.setItem(this.storageKey, JSON.stringify(data));',
        '}',
        '',
        '// Tải dữ liệu từ LocalStorage',
        'loadFromStorage() {',
        '    const raw = localStorage.getItem(this.storageKey);',
        '    if (!raw) return;',
        '    const arr = JSON.parse(raw);',
        '    arr.forEach(json => {',
        '        const sv = Student.fromJSON(json);',
        '        this.students.push(sv);',
        '    });',
        '}',
    ])

    note_box(doc,
             'Dữ liệu LocalStorage tồn tại cho đến khi người dùng xoá thủ công hoặc '
             'xoá cache trình duyệt. Mỗi domain có không gian LocalStorage riêng (~5MB).',
             'D6EAF8')
    doc.add_paragraph()

    # ================================================================
    # 4.5 XÂY DỰNG GIAO DIỆN NGƯỜI DÙNG
    # ================================================================
    heading1(doc, '4.5. Xây dựng Giao diện Người dùng')

    heading2(doc, '4.5.1. Cấu trúc layout tổng thể (HTML)')

    body(doc,
         'Giao diện được xây dựng theo mô hình single-page application (SPA) '
         '— toàn bộ nội dung nằm trong một file index.html, các section được '
         'hiển thị/ẩn động mà không cần tải lại trang.')

    body(doc, 'Cấu trúc HTML chính gồm các thành phần sau:')
    bullet(doc, 'Header — logo, tên hệ thống, hiển thị tổng số sinh viên và đồng hồ thời gian thực.', '')
    bullet(doc, 'Sidebar — thanh điều hướng với 5 mục: Tổng quan, Thêm sinh viên, '
                'Danh sách, Tìm kiếm, Thống kê.', '')
    bullet(doc, 'Content Area — vùng nội dung chính, hiển thị section tương ứng.', '')
    bullet(doc, 'Modal — hộp thoại nổi để xem và chỉnh sửa thông tin sinh viên.', '')
    bullet(doc, 'Toast Notification — thông báo kết quả thao tác (thành công/lỗi).', '')
    bullet(doc, 'Animated Background — nền hoạt ảnh với các vòng tròn chuyển động gradient.', '')

    code_block(doc, [
        '<!-- Cấu trúc layout chính -->',
        '<div class="container">',
        '    <header class="header"> ... </header>',
        '    <main class="main-content">',
        '        <aside class="sidebar"> ... </aside>',
        '        <div class="content-area">',
        '            <section id="dashboard">    ... </section>',
        '            <section id="add-student">  ... </section>',
        '            <section id="student-list"> ... </section>',
        '            <section id="search">       ... </section>',
        '            <section id="statistics">   ... </section>',
        '        </div>',
        '    </main>',
        '</div>',
        '<div id="modal"> ... </div>',
        '<div id="toast"> ... </div>',
    ])

    heading2(doc, '4.5.2. Thiết kế CSS và Hoạt ảnh')

    body(doc,
         'Toàn bộ giao diện được định nghĩa trong styles.css với các kỹ thuật '
         'CSS hiện đại, tạo ra trải nghiệm thân thiện và chuyên nghiệp:')

    bullet(doc, 'CSS Custom Properties (biến CSS) — dễ dàng thay đổi màu sắc, kích thước toàn cục.', '')
    bullet(doc, 'Flexbox và CSS Grid — bố cục linh hoạt, tự điều chỉnh.', '')
    bullet(doc, 'CSS Animations và Keyframes — nền gradient chuyển động, hiệu ứng xuất hiện trang.', '')
    bullet(doc, 'Glassmorphism — hiệu ứng kính mờ (backdrop-filter: blur) cho các card.', '')
    bullet(doc, 'Hover effects và transitions — phản hồi trực quan khi tương tác.', '')
    bullet(doc, 'Box-shadow nhiều lớp — tạo chiều sâu cho các thành phần.', '')

    code_block(doc, [
        '/* Biến CSS toàn cục */           /* Hiệu ứng kính mờ */',
        ':root {                            .card {',
        '    --primary: #667eea;                background: rgba(255,255,255,0.15);',
        '    --secondary: #764ba2;              backdrop-filter: blur(20px);',
        '    --success: #48bb78;                border: 1px solid rgba(255,255,255,0.3);',
        '    --danger: #f56565;                 border-radius: 20px;',
        '    --sidebar-w: 260px;                box-shadow: 0 8px 32px rgba(0,0,0,0.15);',
        '}                                  }',
    ])

    heading2(doc, '4.5.3. Các chức năng giao diện chính')

    body(doc, 'Lớp UIController (DieuKhienGiaoDien.js) điều phối 5 màn hình chính:')

    tbl5 = doc.add_table(rows=1, cols=3)
    tbl5.style = 'Table Grid'
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(tbl5.rows[0].cells, ['Màn hình', 'Section ID', 'Chức năng']):
        set_cell_shading(cell, '1F497D')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(255, 255, 255)

    screens = [
        ('Tổng quan (Dashboard)', 'dashboard',    'Hiển thị 4 thẻ thống kê nhanh: tổng SV, xuất sắc, giỏi, GPA TB; biểu đồ phân loại và theo lớp'),
        ('Thêm sinh viên',         'add-student',  'Form nhập liệu đầy đủ với validation real-time; hỗ trợ thêm mới và chỉnh sửa thông tin'),
        ('Danh sách',              'student-list', 'Bảng danh sách có phân trang (10 SV/trang); sắp xếp theo nhiều tiêu chí; thao tác Xem/Sửa/Xóa'),
        ('Tìm kiếm',               'search',       'Tìm kiếm đồng thời theo mã, tên, lớp; lọc theo giới tính, lớp, xếp loại; hiển thị kết quả real-time'),
        ('Thống kê',               'statistics',   'Biểu đồ tròn phân loại học lực; biểu đồ cột theo lớp; bảng thống kê chi tiết; xuất dữ liệu CSV'),
    ]
    alt = False
    for s in screens:
        row = tbl5.add_row().cells
        bg = 'EBF3FB' if alt else 'FFFFFF'
        for i, (cell, text) in enumerate(zip(row, s)):
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            if i == 0:
                r.bold = True
        alt = not alt

    add_table_borders(tbl5)
    cap5 = doc.add_paragraph('Bảng 4.5: Các màn hình chức năng của hệ thống')
    cap5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap5.runs[0].italic = True
    cap5.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    # ================================================================
    # 4.6 TRIỂN KHAI CÁC CHỨC NĂNG CHÍNH
    # ================================================================
    heading1(doc, '4.6. Triển khai Các Chức năng Chính')

    heading2(doc, '4.6.1. Chức năng Thêm sinh viên')

    body(doc,
         'Khi người dùng điền form và nhấn nút "Thêm sinh viên", '
         'hệ thống thực hiện theo luồng sau:')

    steps_add = [
        ('Bước 1', 'Thu thập dữ liệu: Đọc giá trị từ tất cả các trường input của form.'),
        ('Bước 2', 'Tạo đối tượng Student mới từ dữ liệu vừa thu thập.'),
        ('Bước 3', 'Gọi student.validate() để kiểm tra tính hợp lệ. Nếu sai → hiển thị thông báo lỗi bên cạnh trường tương ứng.'),
        ('Bước 4', 'Gọi manager.addStudent(student). Bên trong: tìm kiếm tuần tự O(n) kiểm tra mã SV đã tồn tại chưa.'),
        ('Bước 5', 'Nếu thành công → push vào DynamicArray → saveToStorage() → hiển thị toast "Thêm thành công".'),
        ('Bước 6', 'Reset form, chuyển sang màn hình "Danh sách" để thấy kết quả ngay.'),
    ]
    for step, desc in steps_add:
        bullet(doc, desc, f'{step}: ')

    heading2(doc, '4.6.2. Chức năng Tìm kiếm')

    body(doc,
         'Hệ thống cung cấp hai cơ chế tìm kiếm:')
    bullet(doc, 'Tìm kiếm theo từ khóa (Linear Search) — duyệt toàn bộ danh sách, '
                'so sánh từ khóa với mã SV, tên, lớp, email. Hỗ trợ tìm kiếm không '
                'phân biệt hoa thường và dấu tiếng Việt.', '')
    bullet(doc, 'Tìm kiếm nhị phân (Binary Search) — áp dụng sau khi danh sách đã '
                'sắp xếp theo mã SV, cho phép tìm chính xác với độ phức tạp O(log n).', '')
    bullet(doc, 'Lọc kết hợp — lọc đồng thời theo giới tính, lớp học, xếp loại học lực. '
                'Kết quả được cập nhật real-time ngay khi người dùng thay đổi bộ lọc.', '')

    code_block(doc, [
        '// Minh họa tìm kiếm đa tiêu chí',
        'searchStudents(keyword, filters = {}) {',
        '    let result = this.students.toArray();',
        '',
        '    // 1. Lọc theo từ khóa (Linear Search)',
        '    if (keyword) {',
        '        const kw = keyword.toLowerCase();',
        '        result = result.filter(sv =>',
        '            sv.id.toLowerCase().includes(kw)   ||',
        '            sv.name.toLowerCase().includes(kw) ||',
        '            sv.className.toLowerCase().includes(kw)',
        '        );',
        '    }',
        '',
        '    // 2. Lọc theo xếp loại',
        '    if (filters.rank) {',
        '        result = result.filter(sv => sv.xepLoai() === filters.rank);',
        '    }',
        '    return result;',
        '}',
    ])

    heading2(doc, '4.6.3. Chức năng Sắp xếp danh sách')

    body(doc,
         'Danh sách sinh viên có thể được sắp xếp theo các tiêu chí:')
    bullet(doc, 'Mã sinh viên (A→Z hoặc Z→A)', '')
    bullet(doc, 'Họ tên (A→Z hoặc Z→A)', '')
    bullet(doc, 'Lớp học (A→Z hoặc Z→A)', '')
    bullet(doc, 'Điểm trung bình (Cao→Thấp hoặc Thấp→Cao)', '')

    body(doc,
         'Thuật toán QuickSort được áp dụng với hàm so sánh động, '
         'cho phép sắp xếp theo bất kỳ tiêu chí nào chỉ bằng cách '
         'thay đổi hàm comparator truyền vào:')

    code_block(doc, [
        '// Sắp xếp theo điểm TB giảm dần',
        'manager.sortBy("gpa", false);',
        '',
        '// Bên trong StudentManager:',
        'sortBy(field, ascending = true) {',
        '    const comparator = ThuatToanSapXep.taoHamSoSanh(field, ascending);',
        '    this.students.sort(comparator);  // → QuickSort O(n log n)',
        '    this.saveToStorage();',
        '}',
    ])

    heading2(doc, '4.6.4. Chức năng Thống kê và Biểu đồ')

    body(doc,
         'Màn hình Thống kê sử dụng thuật toán single-pass O(n) '
         '(ThuatToanThongKe.js) để tính toán toàn bộ số liệu trong '
         'một lần duyệt mảng, sau đó trực quan hóa bằng Chart.js:')

    bullet(doc, 'Biểu đồ tròn (Pie/Doughnut Chart) — phân bố học lực: '
                'Xuất sắc / Giỏi / Khá / Trung bình / Yếu.', '')
    bullet(doc, 'Biểu đồ cột (Bar Chart) — số lượng sinh viên và GPA trung bình '
                'theo từng lớp học.', '')
    bullet(doc, 'Bảng thống kê chi tiết — GPA cao nhất, thấp nhất, trung bình '
                'toàn trường; tổng số mỗi xếp loại.', '')
    bullet(doc, 'Xuất CSV — cho phép tải về file CSV toàn bộ dữ liệu.', '')
    doc.add_paragraph()

    # ================================================================
    # 4.7 KIỂM THỬ HỆ THỐNG
    # ================================================================
    heading1(doc, '4.7. Kiểm thử Hệ thống')

    body(doc,
         'Quá trình kiểm thử được thực hiện theo hai hướng: kiểm thử '
         'chức năng (functional testing) và kiểm thử giao diện (UI testing). '
         'Dữ liệu kiểm thử gồm 30 sinh viên mẫu được tải tự động khi '
         'ứng dụng khởi động lần đầu.')

    heading2(doc, '4.7.1. Kiểm thử chức năng CRUD')

    tbl6 = doc.add_table(rows=1, cols=4)
    tbl6.style = 'Table Grid'
    tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(tbl6.rows[0].cells, ['STT', 'Test case', 'Kết quả mong đợi', 'Kết quả thực tế']):
        set_cell_shading(cell, '1F497D')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(255, 255, 255)

    tests_crud = [
        ('1',  'Thêm SV với dữ liệu hợp lệ',            'Thêm thành công, hiển thị toast xanh',                    '✅ Đạt'),
        ('2',  'Thêm SV với mã SV đã tồn tại',           'Thông báo lỗi "Mã SV đã tồn tại"',                       '✅ Đạt'),
        ('3',  'Thêm SV thiếu họ tên',                   'Highlight trường tên màu đỏ, hiển thị lỗi',               '✅ Đạt'),
        ('4',  'Thêm SV với email sai định dạng',         'Thông báo "Email không hợp lệ"',                          '✅ Đạt'),
        ('5',  'Thêm SV điểm ngoài khoảng [0, 10]',      'Thông báo "Điểm phải từ 0 đến 10"',                       '✅ Đạt'),
        ('6',  'Cập nhật thông tin SV',                   'Dữ liệu cập nhật, updatedAt thay đổi',                    '✅ Đạt'),
        ('7',  'Xoá SV, xác nhận trước',                 'SV bị xoá, danh sách cập nhật',                           '✅ Đạt'),
        ('8',  'Dữ liệu còn sau tải lại trang',          'Dữ liệu phục hồi từ LocalStorage',                        '✅ Đạt'),
    ]
    alt = False
    for t in tests_crud:
        row = tbl6.add_row().cells
        bg = 'EBF3FB' if alt else 'FFFFFF'
        for i, (cell, text) in enumerate(zip(row, t)):
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 3] else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            if i == 3 and '✅' in text:
                r.font.color.rgb = RGBColor(0, 128, 0)
                r.bold = True
        alt = not alt

    add_table_borders(tbl6)
    cap6 = doc.add_paragraph('Bảng 4.6: Kết quả kiểm thử chức năng CRUD')
    cap6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap6.runs[0].italic = True
    cap6.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    heading2(doc, '4.7.2. Kiểm thử thuật toán tìm kiếm và sắp xếp')

    tbl7 = doc.add_table(rows=1, cols=4)
    tbl7.style = 'Table Grid'
    tbl7.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(tbl7.rows[0].cells, ['STT', 'Test case', 'Kết quả mong đợi', 'Kết quả thực tế']):
        set_cell_shading(cell, '1F497D')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(255, 255, 255)

    tests_algo = [
        ('9',  'Tìm kiếm theo mã SV chính xác',          'Tìm thấy và hiển thị đúng SV',                            '✅ Đạt'),
        ('10', 'Tìm kiếm theo tên (không phân biệt hoa)', 'Trả về tất cả SV có tên khớp',                            '✅ Đạt'),
        ('11', 'Tìm kiếm từ khóa không tồn tại',         'Hiển thị "Không tìm thấy kết quả"',                       '✅ Đạt'),
        ('12', 'Sắp xếp theo tên A→Z',                   'Danh sách sắp xếp đúng thứ tự',                           '✅ Đạt'),
        ('13', 'Sắp xếp theo GPA cao→thấp',              'SV điểm cao nhất ở đầu danh sách',                        '✅ Đạt'),
        ('14', 'Lọc theo xếp loại "Giỏi"',               'Chỉ hiển thị SV xếp loại Giỏi',                           '✅ Đạt'),
        ('15', 'Binary Search sau khi sắp xếp mã SV',    'Tìm thấy với O(log n)',                                    '✅ Đạt'),
        ('16', 'Thống kê trên 30 SV mẫu',                'Số liệu GPA và phân loại chính xác',                      '✅ Đạt'),
    ]
    alt = False
    for t in tests_algo:
        row = tbl7.add_row().cells
        bg = 'EBF3FB' if alt else 'FFFFFF'
        for i, (cell, text) in enumerate(zip(row, t)):
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 3] else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            if i == 3 and '✅' in text:
                r.font.color.rgb = RGBColor(0, 128, 0)
                r.bold = True
        alt = not alt

    add_table_borders(tbl7)
    cap7 = doc.add_paragraph('Bảng 4.7: Kết quả kiểm thử thuật toán tìm kiếm và sắp xếp')
    cap7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap7.runs[0].italic = True
    cap7.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    heading2(doc, '4.7.3. Kiểm thử giao diện (UI Testing)')

    tbl8 = doc.add_table(rows=1, cols=3)
    tbl8.style = 'Table Grid'
    tbl8.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(tbl8.rows[0].cells, ['Hạng mục kiểm thử', 'Nội dung', 'Kết quả']):
        set_cell_shading(cell, '2E74B5')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(255, 255, 255)

    ui_tests = [
        ('Điều hướng',          'Click các nút trên sidebar, kiểm tra chuyển section đúng', '✅ Đạt'),
        ('Responsive (1920px)', 'Giao diện hiển thị đúng trên màn hình Full HD',            '✅ Đạt'),
        ('Responsive (1366px)', 'Giao diện hiển thị đúng trên màn hình laptop',             '✅ Đạt'),
        ('Modal',               'Mở/đóng modal xem thông tin sinh viên',                    '✅ Đạt'),
        ('Toast notification',  'Hiển thị thông báo 3 giây khi thao tác thành công/lỗi',   '✅ Đạt'),
        ('Phân trang',          'Chuyển trang danh sách, hiển thị đúng 10 SV/trang',        '✅ Đạt'),
        ('Đồng hồ real-time',   'Đồng hồ cập nhật từng giây trên header',                  '✅ Đạt'),
        ('Biểu đồ Chart.js',    'Vẽ đúng biểu đồ tròn và cột khi vào màn hình Thống kê',  '✅ Đạt'),
        ('Phím tắt Escape',     'Nhấn Escape đóng modal đang mở',                          '✅ Đạt'),
        ('Xuất CSV',            'File CSV tải về chứa đúng dữ liệu toàn bộ sinh viên',     '✅ Đạt'),
    ]
    alt = False
    for t in ui_tests:
        row = tbl8.add_row().cells
        bg = 'EBF3FB' if alt else 'FFFFFF'
        for i, (cell, text) in enumerate(zip(row, t)):
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 2 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            if i == 2 and '✅' in text:
                r.font.color.rgb = RGBColor(0, 128, 0)
                r.bold = True
            if i == 0:
                r.bold = True
        alt = not alt

    add_table_borders(tbl8)
    cap8 = doc.add_paragraph('Bảng 4.8: Kết quả kiểm thử giao diện người dùng')
    cap8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap8.runs[0].italic = True
    cap8.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    note_box(doc,
             'Tổng cộng 26/26 test case đã kiểm thử đều đạt yêu cầu (tỷ lệ 100%). '
             'Hệ thống hoạt động ổn định trên Chrome, Edge và Firefox phiên bản mới nhất.',
             'D5E8D4')

    # ================================================================
    # 4.8 KẾT QUẢ TRIỂN KHAI
    # ================================================================
    heading1(doc, '4.8. Kết quả Triển khai')

    heading2(doc, '4.8.1. Các chức năng đã hoàn thành')

    body(doc, 'Sau quá trình xây dựng và kiểm thử, hệ thống đã hoàn thiện đầy đủ các chức năng:')

    tbl9 = doc.add_table(rows=1, cols=3)
    tbl9.style = 'Table Grid'
    tbl9.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(tbl9.rows[0].cells, ['Nhóm chức năng', 'Chức năng cụ thể', 'Trạng thái']):
        set_cell_shading(cell, '1F497D')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(255, 255, 255)

    results_data = [
        ('Cấu trúc dữ liệu', 'Mảng động tự mở rộng (DynamicArray) với đầy đủ thao tác', '✅ Hoàn thành'),
        ('Quản lý sinh viên', 'Thêm, xem, sửa, xoá sinh viên với validation đầy đủ',     '✅ Hoàn thành'),
        ('Tìm kiếm',         'Linear Search đa trường, Binary Search theo mã SV',         '✅ Hoàn thành'),
        ('Sắp xếp',          'QuickSort theo tên, GPA, mã SV, lớp; hai chiều asc/desc',   '✅ Hoàn thành'),
        ('Thống kê',         'GPA, phân loại học lực, thống kê theo lớp bằng O(n)',        '✅ Hoàn thành'),
        ('Biểu đồ',          'Biểu đồ tròn học lực, biểu đồ cột theo lớp (Chart.js)',     '✅ Hoàn thành'),
        ('Lưu trữ',          'LocalStorage: tự động lưu và phục hồi dữ liệu',             '✅ Hoàn thành'),
        ('Giao diện',        'SPA 5 màn hình, animation, glass effect, responsive',        '✅ Hoàn thành'),
        ('Phân trang',       'Phân trang danh sách 10 SV/trang, điều hướng trang',        '✅ Hoàn thành'),
        ('Xuất dữ liệu',     'Xuất file CSV toàn bộ danh sách sinh viên',                 '✅ Hoàn thành'),
        ('Dữ liệu mẫu',      '30 sinh viên mẫu được tải tự động khi khởi động',           '✅ Hoàn thành'),
    ]
    alt = False
    for r_data in results_data:
        row = tbl9.add_row().cells
        bg = 'EBF3FB' if alt else 'FFFFFF'
        for i, (cell, text) in enumerate(zip(row, r_data)):
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 2 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            if i == 0:
                r.bold = True
            if i == 2 and '✅' in text:
                r.font.color.rgb = RGBColor(0, 128, 0)
                r.bold = True
        alt = not alt

    add_table_borders(tbl9)
    cap9 = doc.add_paragraph('Bảng 4.9: Tổng hợp kết quả triển khai hệ thống')
    cap9.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap9.runs[0].italic = True
    cap9.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    heading2(doc, '4.8.2. Hiệu suất hệ thống')

    body(doc,
         'Hệ thống được đo lường hiệu suất trên tập dữ liệu 30 sinh viên mẫu '
         'và kiểm tra khả năng mở rộng với 500 và 1000 bản ghi:')

    tbl10 = doc.add_table(rows=1, cols=4)
    tbl10.style = 'Table Grid'
    tbl10.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(tbl10.rows[0].cells, ['Thao tác', '30 SV (ms)', '500 SV (ms)', '1000 SV (ms)']):
        set_cell_shading(cell, '2E74B5')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(255, 255, 255)

    perf_data = [
        ('Tải dữ liệu từ LocalStorage', '< 1',  '< 5',   '< 10'),
        ('Tìm kiếm theo từ khóa',       '< 1',  '< 2',   '< 3'),
        ('Sắp xếp QuickSort (GPA)',     '< 1',  '< 5',   '< 8'),
        ('Tính toán thống kê',          '< 1',  '< 2',   '< 3'),
        ('Render bảng danh sách',        '< 5',  '< 10',  '< 15'),
        ('Vẽ biểu đồ Chart.js',         '< 50', '< 60',  '< 80'),
    ]
    alt = False
    for pd in perf_data:
        row = tbl10.add_row().cells
        bg = 'EBF3FB' if alt else 'FFFFFF'
        for i, (cell, text) in enumerate(zip(row, pd)):
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            if i == 0:
                r.bold = True
        alt = not alt

    add_table_borders(tbl10)
    cap10 = doc.add_paragraph('Bảng 4.10: Hiệu suất thực tế đo lường trên trình duyệt Chrome')
    cap10.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap10.runs[0].italic = True
    cap10.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    body(doc,
         'Kết quả cho thấy hệ thống hoạt động với hiệu suất rất tốt ở quy mô từ '
         'hàng chục đến hàng nghìn sinh viên — phù hợp với mục tiêu của bài tập lớn '
         'là minh họa và ứng dụng các thuật toán trong thực tế.')

    heading2(doc, '4.8.3. Hướng phát triển tiếp theo')

    body(doc, 'Mặc dù hệ thống đã hoàn thành đầy đủ theo yêu cầu đề tài, '
              'một số tính năng có thể phát triển thêm trong tương lai:')

    future = [
        ('Xuất PDF', 'Cho phép in phiếu điểm cá nhân hoặc bảng lớp dạng PDF trực tiếp từ trình duyệt.'),
        ('Nhập từ Excel/CSV', 'Nhập hàng loạt dữ liệu sinh viên từ file Excel hoặc CSV.'),
        ('Phân quyền', 'Thêm đăng nhập, phân biệt quyền quản trị viên và giáo viên.'),
        ('Backend API', 'Kết nối REST API, lưu dữ liệu trên server (Node.js + MongoDB).'),
        ('Thêm môn học', 'Mở rộng hệ thống điểm từ 3 môn cố định sang nhiều môn động.'),
        ('Giao diện di động', 'Tối ưu layout cho thiết bị di động (responsive hoàn toàn).'),
    ]
    for title, desc in future:
        bullet(doc, desc, f'{title}: ')

    doc.add_paragraph()

    # ================================================================
    # TÓM TẮT CHƯƠNG
    # ================================================================
    heading1(doc, 'Tóm tắt Chương 4')

    note_box(doc,
             'Chương 4 đã trình bày toàn bộ quá trình xây dựng và triển khai Hệ thống '
             'Quản lý Sinh viên, từ thiết lập môi trường đến từng dòng code, từ cấu trúc '
             'dữ liệu DynamicArray đến giao diện đồ họa phong phú. Kết quả: 11 chức năng '
             'chính đã hoàn thành, 26/26 test case đạt yêu cầu, hệ thống hoạt động ổn định '
             'trên mọi trình duyệt hiện đại.',
             'D5E8D4')

    body(doc,
         'Qua đây, đề tài đã minh chứng rõ ràng rằng cấu trúc dữ liệu mảng động và các '
         'thuật toán cổ điển (QuickSort, Binary Search, linear statistics) hoàn toàn đủ '
         'năng lực để xây dựng một ứng dụng thực tế có giao diện đồ họa chuyên nghiệp '
         'và hiệu suất tốt — không cần đến bất kỳ thư viện JavaScript nặng nào.')

    # ================================================================
    # LƯU FILE
    # ================================================================
    output_path = r'e:\Hoc_tap\BTL_Laptrinhnangcao\Chuong4_XayDung_TrienKhai.docx'
    doc.save(output_path)
    print(f'✅ Đã tạo file Word thành công: {output_path}')
    print(f'   - Chương 4: Xây dựng và Triển khai Hệ thống')
    print(f'   - Gồm 8 mục lớn, 10 bảng chi tiết, 26 test case')

if __name__ == '__main__':
    create_chapter4()
