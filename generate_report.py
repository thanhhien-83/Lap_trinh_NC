# -*- coding: utf-8 -*-
"""
Tạo báo cáo Word chi tiết cho BTL Lập trình nâng cao
Đề tài: Xây dựng phần mềm quản lý sinh viên sử dụng cấu trúc mảng động và giao diện đồ họa
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Tô màu nền cho ô trong bảng"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_borders(table):
    """Thêm viền cho bảng"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)

def create_report():
    doc = Document()
    
    # ======================== THIẾT LẬP STYLE ========================
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    # Thiết lập font cho các heading
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0, 0, 139)
        if i == 1:
            heading_style.font.size = Pt(18)
        elif i == 2:
            heading_style.font.size = Pt(15)
        else:
            heading_style.font.size = Pt(13)
    
    # ======================== TRANG BÌA ========================
    for _ in range(4):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('BÁO CÁO BÀI TẬP LỚN')
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 139)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('LẬP TRÌNH NÂNG CAO')
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 139)
    
    doc.add_paragraph()
    
    topic = doc.add_paragraph()
    topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = topic.add_run('Đề tài:')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    
    topic2 = doc.add_paragraph()
    topic2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = topic2.add_run('"Xây dựng phần mềm quản lý sinh viên\nsử dụng cấu trúc mảng động và giao diện đồ họa"')
    run.bold = True
    run.italic = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(139, 0, 0)
    
    for _ in range(4):
        doc.add_paragraph()
    
    # Ngôn ngữ & Công nghệ
    tech = doc.add_paragraph()
    tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tech.add_run('Ngôn ngữ: JavaScript | Giao diện: HTML/CSS\nCấu trúc dữ liệu: Mảng động (Dynamic Array)')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.italic = True
    
    doc.add_page_break()
    
    # ======================== MỤC LỤC ========================
    doc.add_heading('MỤC LỤC', level=1)
    
    toc_items = [
        ('CHƯƠNG 1: CƠ SỞ LÝ THUYẾT', True, 0),
        ('1.1. Tổng quan về cấu trúc dữ liệu', False, 1),
        ('1.2. Mảng động (Dynamic Array)', False, 1),
        ('1.3. Lập trình hướng đối tượng (OOP)', False, 1),
        ('1.4. Thuật toán sắp xếp', False, 1),
        ('1.5. Thuật toán tìm kiếm', False, 1),
        ('1.6. Thuật toán thống kê', False, 1),
        ('1.7. Phân tích độ phức tạp thuật toán', False, 1),
        ('CHƯƠNG 2: PHÂN TÍCH BÀI TOÁN', True, 0),
        ('2.1. Mô tả bài toán', False, 1),
        ('2.2. Yêu cầu chức năng', False, 1),
        ('2.3. Kiến trúc hệ thống', False, 1),
        ('2.4. Mô hình đối tượng', False, 1),
        ('2.5. Thiết kế cấu trúc dữ liệu', False, 1),
        ('2.6. Luồng xử lý chính', False, 1),
        ('CHƯƠNG 3: THUẬT TOÁN SỬ DỤNG', True, 0),
        ('3.1. Thuật toán trên mảng động', False, 1),
        ('3.2. Thuật toán sắp xếp QuickSort', False, 1),
        ('3.3. Thuật toán sắp xếp BubbleSort', False, 1),
        ('3.4. Thuật toán tìm kiếm tuần tự', False, 1),
        ('3.5. Thuật toán tìm kiếm nhị phân', False, 1),
        ('3.6. Thuật toán thống kê Single-pass', False, 1),
        ('3.7. Thuật toán phân loại học lực', False, 1),
        ('3.8. Thuật toán Validation dữ liệu', False, 1),
        ('3.9. Bảng tổng hợp độ phức tạp', False, 1),
    ]
    
    for text, is_bold, indent_level in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if indent_level > 0:
            p.paragraph_format.left_indent = Cm(1.5)
        run = p.add_run(text)
        run.bold = is_bold
        run.font.size = Pt(13)
        run.font.name = 'Times New Roman'
    
    doc.add_page_break()
    
    # ================================================================
    #                    CHƯƠNG 1: CƠ SỞ LÝ THUYẾT
    # ================================================================
    doc.add_heading('CHƯƠNG 1: CƠ SỞ LÝ THUYẾT', level=1)
    
    # --- 1.1 ---
    doc.add_heading('1.1. Tổng quan về cấu trúc dữ liệu', level=2)
    
    doc.add_paragraph(
        'Cấu trúc dữ liệu (Data Structure) là cách thức tổ chức, quản lý và lưu trữ dữ liệu trong máy tính sao cho '
        'việc truy cập và thao tác trên dữ liệu được thực hiện một cách hiệu quả nhất. Việc lựa chọn cấu trúc dữ liệu '
        'phù hợp có ảnh hưởng trực tiếp đến hiệu suất của chương trình, bao gồm tốc độ thực thi, lượng bộ nhớ tiêu thụ '
        'và khả năng mở rộng của hệ thống.'
    )
    
    doc.add_paragraph(
        'Trong lĩnh vực khoa học máy tính, cấu trúc dữ liệu được phân thành nhiều loại khác nhau, mỗi loại có những '
        'ưu điểm và nhược điểm riêng. Các cấu trúc dữ liệu cơ bản bao gồm: mảng (array), danh sách liên kết (linked list), '
        'ngăn xếp (stack), hàng đợi (queue), cây (tree), đồ thị (graph) và bảng băm (hash table). Mỗi cấu trúc dữ liệu '
        'được thiết kế để giải quyết một nhóm bài toán cụ thể và có độ phức tạp thời gian khác nhau cho các thao tác cơ bản '
        'như thêm, xóa, tìm kiếm và truy cập phần tử.'
    )
    
    doc.add_paragraph(
        'Trong bài tập lớn này, chúng tôi lựa chọn cấu trúc dữ liệu Mảng động (Dynamic Array) làm nền tảng chính để '
        'xây dựng hệ thống quản lý sinh viên. Mảng động là một phiên bản cải tiến của mảng tĩnh truyền thống, cho phép '
        'tự động mở rộng kích thước khi cần thiết, đồng thời vẫn giữ được ưu điểm truy cập trực tiếp O(1) của mảng.'
    )
    
    # --- 1.2 ---
    doc.add_heading('1.2. Mảng động (Dynamic Array)', level=2)
    
    doc.add_heading('1.2.1. Khái niệm', level=3)
    doc.add_paragraph(
        'Mảng tĩnh (Static Array) là cấu trúc dữ liệu cơ bản nhất, trong đó các phần tử được lưu trữ liên tiếp nhau '
        'trong bộ nhớ với kích thước cố định ngay từ khi khai báo. Ví dụ, khi khai báo một mảng tĩnh có 100 phần tử, '
        'chương trình sẽ cấp phát đúng 100 ô nhớ liên tiếp, và kích thước này không thể thay đổi trong suốt quá trình '
        'chạy chương trình. Điều này dẫn đến hai vấn đề lớn: (1) Nếu số phần tử cần lưu trữ vượt quá 100, chương trình '
        'sẽ bị tràn mảng (overflow); (2) Nếu chỉ sử dụng 10 ô nhớ, 90 ô nhớ còn lại bị lãng phí.'
    )
    
    doc.add_paragraph(
        'Mảng động (Dynamic Array) ra đời để khắc phục những hạn chế trên. Mảng động có khả năng tự động mở rộng kích thước '
        'khi mảng đầy. Cơ chế hoạt động cơ bản của mảng động như sau: khởi tạo với một dung lượng ban đầu nhỏ (ví dụ 10 phần tử), '
        'khi thêm phần tử mà mảng đã đầy (length ≥ capacity), hệ thống sẽ tự động tạo một mảng mới có dung lượng gấp đôi, '
        'sao chép toàn bộ dữ liệu từ mảng cũ sang mảng mới, rồi tiếp tục sử dụng mảng mới. Quá trình này được gọi là '
        'cấp phát lại (resize).'
    )
    
    doc.add_heading('1.2.2. Minh họa hoạt động mảng động', level=3)
    doc.add_paragraph(
        'Để hiểu rõ hơn cách hoạt động của mảng động, chúng ta hãy xem xét ví dụ sau với mảng có dung lượng ban đầu là 4:'
    )
    
    p = doc.add_paragraph()
    run = p.add_run(
        'Bước 1: Khởi tạo mảng — capacity = 4, length = 0\n'
        '    [_][_][_][_]\n\n'
        'Bước 2: Thêm 3 phần tử (A, B, C) — length = 3\n'
        '    [A][B][C][_]\n\n'
        'Bước 3: Thêm phần tử thứ 4 (D) — length = 4 (ĐẦY!)\n'
        '    [A][B][C][D]\n\n'
        'Bước 4: Thêm phần tử thứ 5 (E) → cần resize:\n'
        '    ► Tạo mảng mới capacity = 8\n'
        '    ► Copy dữ liệu cũ sang mảng mới\n'
        '    ► Thêm phần tử E\n'
        '    [A][B][C][D][E][_][_][_]'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_heading('1.2.3. Phân tích khấu hao (Amortized Analysis)', level=3)
    doc.add_paragraph(
        'Một câu hỏi quan trọng khi sử dụng mảng động là: tại sao lại gấp đôi dung lượng khi resize? '
        'Câu trả lời nằm ở phân tích khấu hao (amortized analysis) — một phương pháp đánh giá chi phí trung bình '
        'của một chuỗi thao tác, thay vì chỉ xem xét chi phí của từng thao tác riêng lẻ.'
    )
    
    doc.add_paragraph(
        'Nếu mỗi lần resize chỉ tăng thêm 1 ô nhớ (tức capacity = capacity + 1), thì mỗi lần thêm phần tử khi mảng đầy '
        'đều phải resize, tức phải sao chép toàn bộ n phần tử → chi phí O(n) cho MỖI lần push. Tổng chi phí cho n lần push '
        'sẽ là O(n²), cực kỳ kém hiệu quả.'
    )
    
    doc.add_paragraph(
        'Ngược lại, nếu gấp đôi dung lượng (capacity = capacity × 2), thì sau k lần resize, tổng chi phí sao chép là: '
        'n + n/2 + n/4 + n/8 + ... ≈ 2n. Chia đều cho n lần push, chi phí trung bình mỗi lần push chỉ là O(2n/n) = O(1). '
        'Đây chính là nguyên lý phân tích khấu hao: mặc dù một vài thao tác riêng lẻ có thể rất tốn kém (O(n) khi resize), '
        'nhưng tính trung bình trên toàn bộ chuỗi thao tác, chi phí mỗi thao tác chỉ là O(1).'
    )

    doc.add_heading('1.2.4. Ba thuộc tính cốt lõi', level=3)
    doc.add_paragraph(
        'Trong hệ thống của chúng tôi, mảng động được triển khai với ba thuộc tính cốt lõi:'
    )
    doc.add_paragraph('data: Mảng JavaScript thực tế dùng để lưu trữ các phần tử. Đây là vùng nhớ chính chứa dữ liệu.', style='List Bullet')
    doc.add_paragraph('length: Số lượng phần tử đang thực sự được sử dụng trong mảng (0 ≤ length ≤ capacity).', style='List Bullet')
    doc.add_paragraph('capacity: Dung lượng tối đa hiện tại của mảng. Khi length đạt tới capacity, mảng sẽ được mở rộng.', style='List Bullet')
    
    doc.add_paragraph(
        'Sự phân biệt giữa length và capacity là đặc trưng quan trọng nhất của mảng động. Length cho biết bao nhiêu phần tử '
        'đang được sử dụng, còn capacity cho biết bao nhiêu phần tử có thể chứa trước khi cần resize. Điều này cho phép '
        'mảng động hoạt động hiệu quả: hầu hết các thao tác push đều có chi phí O(1) vì chỉ cần gán giá trị vào vị trí '
        'tiếp theo, chỉ thỉnh thoảng mới phải resize.'
    )
    
    # --- 1.3 ---
    doc.add_heading('1.3. Lập trình hướng đối tượng (OOP)', level=2)
    
    doc.add_paragraph(
        'Lập trình hướng đối tượng (Object-Oriented Programming — OOP) là một phương pháp lập trình dựa trên khái niệm '
        '"đối tượng" (object). Mỗi đối tượng chứa dữ liệu dưới dạng các thuộc tính (attributes/properties) và hành vi '
        'dưới dạng các phương thức (methods). OOP giúp tổ chức mã nguồn rõ ràng, dễ bảo trì và mở rộng thông qua bốn '
        'nguyên lý cơ bản: đóng gói (encapsulation), kế thừa (inheritance), đa hình (polymorphism) và trừu tượng hóa (abstraction).'
    )
    
    doc.add_paragraph(
        'Trong hệ thống quản lý sinh viên này, chúng tôi áp dụng OOP với các lớp chính sau:'
    )
    
    doc.add_paragraph('Lớp Student (SinhVien.js): Đại diện cho một sinh viên với các thuộc tính như mã sinh viên, họ tên, ngày sinh, '
                      'giới tính, lớp, email, số điện thoại, địa chỉ và điểm số các môn. Lớp này cũng chứa các phương thức tính '
                      'điểm trung bình, xếp loại học lực và kiểm tra dữ liệu hợp lệ.', style='List Bullet')
    doc.add_paragraph('Lớp DynamicArray (MangDong.js): Triển khai cấu trúc dữ liệu mảng động với các thao tác push, pop, get, set, '
                      'insertAt, removeAt, find, filter, sort, binarySearch và clone.', style='List Bullet')
    doc.add_paragraph('Lớp StudentManager (QuanLySinhVien.js): Quản lý danh sách sinh viên sử dụng mảng động, cung cấp các chức năng '
                      'CRUD (Create, Read, Update, Delete), tìm kiếm, sắp xếp và thống kê.', style='List Bullet')
    doc.add_paragraph('Lớp UIController (DieuKhienGiaoDien.js): Điều khiển giao diện người dùng, xử lý sự kiện, hiển thị bảng biểu '
                      'và biểu đồ thống kê.', style='List Bullet')
    doc.add_paragraph('Lớp Application (UngDung.js): Điểm khởi đầu của ứng dụng, kết nối tất cả các module lại với nhau.', style='List Bullet')
    
    # --- 1.4 ---
    doc.add_heading('1.4. Thuật toán sắp xếp', level=2)
    
    doc.add_heading('1.4.1. Khái niệm', level=3)
    doc.add_paragraph(
        'Sắp xếp (Sorting) là quá trình sắp xếp lại các phần tử trong một tập dữ liệu theo một thứ tự nhất định '
        '(tăng dần hoặc giảm dần) dựa trên một khóa so sánh (sort key). Sắp xếp là một trong những bài toán cơ bản '
        'và quan trọng nhất trong khoa học máy tính, bởi vì nhiều thuật toán khác (như tìm kiếm nhị phân) yêu cầu '
        'dữ liệu đầu vào phải được sắp xếp trước.'
    )
    
    doc.add_heading('1.4.2. QuickSort — Sắp xếp nhanh', level=3)
    doc.add_paragraph(
        'QuickSort là thuật toán sắp xếp dựa trên chiến lược "chia để trị" (Divide and Conquer), được phát minh bởi '
        'Tony Hoare vào năm 1960. Ý tưởng cốt lõi của QuickSort là: chọn một phần tử làm "chốt" (pivot), sau đó '
        'phân hoạch mảng thành hai phần — phần bên trái chứa các phần tử nhỏ hơn hoặc bằng pivot, phần bên phải chứa '
        'các phần tử lớn hơn pivot. Quá trình này được lặp lại đệ quy cho từng phần cho đến khi mảng được sắp xếp hoàn toàn.'
    )
    
    doc.add_paragraph(
        'Các bước thực hiện QuickSort:\n'
        '• Bước 1: Chọn phần tử chốt (pivot). Trong triển khai của chúng tôi, chúng tôi chọn phần tử cuối cùng làm pivot.\n'
        '• Bước 2: Phân hoạch (Partition) — duyệt mảng từ trái sang phải, đưa các phần tử nhỏ hơn pivot về bên trái, '
        'các phần tử lớn hơn pivot về bên phải.\n'
        '• Bước 3: Đặt pivot vào đúng vị trí cuối cùng của nó trong mảng đã sắp xếp.\n'
        '• Bước 4: Đệ quy sắp xếp nửa trái (từ đầu đến pivot - 1).\n'
        '• Bước 5: Đệ quy sắp xếp nửa phải (từ pivot + 1 đến cuối).'
    )
    
    doc.add_paragraph('Độ phức tạp thời gian:')
    doc.add_paragraph('Trường hợp tốt nhất và trung bình: O(n log n) — khi pivot chia đều mảng thành hai nửa xấp xỉ bằng nhau.', style='List Bullet')
    doc.add_paragraph('Trường hợp xấu nhất: O(n²) — khi mảng đã sắp xếp sẵn và luôn chọn phần tử nhỏ nhất/lớn nhất làm pivot.', style='List Bullet')
    doc.add_paragraph('Bộ nhớ phụ: O(log n) — cho ngăn xếp đệ quy.', style='List Bullet')
    
    doc.add_heading('1.4.3. BubbleSort — Sắp xếp nổi bọt', level=3)
    doc.add_paragraph(
        'BubbleSort là một trong những thuật toán sắp xếp đơn giản nhất. Ý tưởng của thuật toán là duyệt mảng nhiều lần, '
        'mỗi lần so sánh hai phần tử liền kề và hoán đổi nếu chúng không đúng thứ tự. Sau mỗi vòng lặp, phần tử lớn nhất '
        '(chưa đúng vị trí) sẽ "nổi" lên cuối mảng, giống như bong bóng nổi lên mặt nước — đó là lý do tên gọi "sắp xếp nổi bọt".'
    )
    
    doc.add_paragraph('Độ phức tạp thời gian:')
    doc.add_paragraph('Trường hợp tốt nhất: O(n) — khi mảng đã sắp xếp sẵn (có cờ dừng sớm).', style='List Bullet')
    doc.add_paragraph('Trường hợp trung bình và xấu nhất: O(n²).', style='List Bullet')
    doc.add_paragraph('Bộ nhớ phụ: O(1) — sắp xếp tại chỗ.', style='List Bullet')
    
    # --- 1.5 ---
    doc.add_heading('1.5. Thuật toán tìm kiếm', level=2)
    
    doc.add_heading('1.5.1. Tìm kiếm tuần tự (Linear Search)', level=3)
    doc.add_paragraph(
        'Tìm kiếm tuần tự là thuật toán tìm kiếm cơ bản nhất: duyệt lần lượt từng phần tử trong mảng, so sánh từng '
        'phần tử với giá trị cần tìm. Nếu tìm thấy phần tử phù hợp, thuật toán trả về vị trí (index) của phần tử đó; '
        'nếu duyệt hết mảng mà không tìm thấy, thuật toán trả về -1 để báo không tìm thấy.'
    )
    
    doc.add_paragraph(
        'Ưu điểm của tìm kiếm tuần tự là đơn giản, dễ cài đặt, và đặc biệt không yêu cầu mảng phải được sắp xếp trước. '
        'Nhược điểm là hiệu suất thấp khi mảng có kích thước lớn (O(n) cho mọi trường hợp).'
    )
    
    doc.add_heading('1.5.2. Tìm kiếm nhị phân (Binary Search)', level=3)
    doc.add_paragraph(
        'Tìm kiếm nhị phân là thuật toán tìm kiếm hiệu quả cao, hoạt động trên mảng ĐÃ ĐƯỢC SẮP XẾP. Ý tưởng của thuật toán '
        'là chia đôi không gian tìm kiếm ở mỗi bước: so sánh phần tử ở giữa mảng với giá trị cần tìm, nếu bằng nhau thì '
        'tìm thấy, nếu nhỏ hơn thì tìm trong nửa phải, nếu lớn hơn thì tìm trong nửa trái. Quá trình lặp lại cho đến khi '
        'tìm thấy hoặc đoạn tìm kiếm rỗng.'
    )
    
    doc.add_paragraph(
        'So sánh hiệu suất với tìm kiếm tuần tự: với n = 1.000 phần tử, tìm kiếm tuần tự cần tối đa 1.000 bước, trong khi '
        'tìm kiếm nhị phân chỉ cần tối đa 10 bước (log₂(1000) ≈ 10). Với n = 1.000.000 phần tử, tìm kiếm tuần tự cần tối đa '
        '1.000.000 bước, còn tìm kiếm nhị phân chỉ cần khoảng 20 bước. Sự khác biệt là cực kỳ lớn.'
    )
    
    # --- 1.6 ---
    doc.add_heading('1.6. Thuật toán thống kê', level=2)
    doc.add_paragraph(
        'Thuật toán thống kê trong hệ thống được thiết kế theo nguyên lý "duyệt một lần duy nhất" (single-pass), nghĩa là '
        'chỉ cần duyệt qua mảng dữ liệu một lần nhưng thu thập được đồng thời nhiều chỉ số thống kê. Trong mỗi bước duyệt, '
        'thuật toán cập nhật song song: tổng điểm (để tính trung bình), điểm cao nhất (max), điểm thấp nhất (min), số lượng '
        'sinh viên theo từng xếp loại, và số lượng sinh viên theo từng lớp.'
    )
    
    doc.add_paragraph(
        'Ưu điểm của phương pháp single-pass: hiệu quả gấp nhiều lần so với việc duyệt riêng lẻ cho từng chỉ số. Nếu cần '
        'tính 5 chỉ số khác nhau và mỗi chỉ số duyệt 1 lần → phải duyệt 5n phần tử. Với single-pass, chỉ cần duyệt n phần tử '
        'nhưng vẫn thu thập được tất cả 5 chỉ số. Thuật toán đếm và phân nhóm sử dụng bảng băm (hash map/object) để đếm '
        'số phần tử thuộc mỗi nhóm, với chi phí O(n) thời gian và O(k) bộ nhớ (k là số nhóm).'
    )
    
    # --- 1.7 ---
    doc.add_heading('1.7. Phân tích độ phức tạp thuật toán', level=2)
    doc.add_paragraph(
        'Độ phức tạp thuật toán (Algorithm Complexity) là thước đo đánh giá hiệu suất của thuật toán, thường được biểu diễn '
        'bằng ký hiệu Big-O. Ký hiệu Big-O mô tả tốc độ tăng trưởng của thời gian thực thi (hoặc bộ nhớ sử dụng) theo '
        'kích thước đầu vào n. Các lớp độ phức tạp phổ biến từ tốt đến tệ: O(1) — hằng số, O(log n) — logarit, O(n) — tuyến tính, '
        'O(n log n) — tuyến tính-logarit, O(n²) — bình phương, O(2ⁿ) — mũ.'
    )
    
    # Bảng Big-O
    table = doc.add_table(rows=7, cols=3)
    add_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Ký hiệu Big-O', 'Tên gọi', 'Ví dụ']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')
    
    data = [
        ['O(1)', 'Hằng số', 'Truy cập mảng theo chỉ số'],
        ['O(log n)', 'Logarit', 'Tìm kiếm nhị phân'],
        ['O(n)', 'Tuyến tính', 'Tìm kiếm tuần tự, duyệt mảng'],
        ['O(n log n)', 'Tuyến tính-logarit', 'QuickSort, MergeSort'],
        ['O(n²)', 'Bình phương', 'BubbleSort, hai vòng lặp lồng nhau'],
        ['O(2ⁿ)', 'Mũ', 'Vét cạn tất cả tập con'],
    ]
    
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = cell_text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
                    r.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ================================================================
    #                    CHƯƠNG 2: PHÂN TÍCH BÀI TOÁN
    # ================================================================
    doc.add_heading('CHƯƠNG 2: PHÂN TÍCH BÀI TOÁN', level=1)
    
    # --- 2.1 ---
    doc.add_heading('2.1. Mô tả bài toán', level=2)
    doc.add_paragraph(
        'Bài toán đặt ra là xây dựng một phần mềm quản lý sinh viên hoàn chỉnh sử dụng cấu trúc dữ liệu mảng động '
        'và giao diện đồ họa. Phần mềm cần đáp ứng các yêu cầu nghiệp vụ thực tế của việc quản lý thông tin sinh viên '
        'tại một trường đại học hoặc cơ sở giáo dục, bao gồm: quản lý thông tin cá nhân, điểm số, xếp loại học lực, '
        'tìm kiếm và thống kê.'
    )
    
    doc.add_paragraph(
        'Mỗi sinh viên trong hệ thống được đặc trưng bởi các thông tin sau: mã sinh viên (duy nhất, dùng để định danh), '
        'họ và tên, ngày sinh, giới tính (Nam/Nữ), lớp, email, số điện thoại, địa chỉ, và điểm số ba môn học (Toán, Văn, '
        'Tiếng Anh). Từ điểm ba môn này, hệ thống sẽ tự động tính toán điểm trung bình và xếp loại học lực theo thang '
        'đánh giá 5 mức: Xuất sắc (≥ 9.0), Giỏi (≥ 8.0), Khá (≥ 6.5), Trung bình (≥ 5.0) và Yếu (< 5.0).'
    )
    
    doc.add_paragraph(
        'Hệ thống được phát triển bằng ngôn ngữ JavaScript thuần, kết hợp HTML và CSS để xây dựng giao diện đồ họa '
        'web. Toàn bộ dữ liệu được quản lý trên cấu trúc mảng động tự viết (không sử dụng mảng JavaScript tích hợp '
        'cho cấu trúc dữ liệu chính), nhằm thể hiện sự hiểu biết sâu sắc về cấu trúc dữ liệu và thuật toán.'
    )
    
    # --- 2.2 ---
    doc.add_heading('2.2. Yêu cầu chức năng', level=2)
    
    doc.add_heading('2.2.1. Quản lý sinh viên (CRUD)', level=3)
    doc.add_paragraph(
        'Hệ thống cần cung cấp đầy đủ bốn thao tác CRUD trên danh sách sinh viên:'
    )
    doc.add_paragraph(
        'Create (Thêm): Cho phép người dùng nhập thông tin sinh viên mới thông qua form giao diện. '
        'Hệ thống phải kiểm tra tính hợp lệ của dữ liệu (validation) trước khi thêm: mã sinh viên không được trùng lặp, '
        'họ tên không được để trống, ngày sinh không được ở tương lai, giới tính phải hợp lệ, điểm số phải nằm trong '
        'khoảng [0, 10]. Thuật toán kiểm tra trùng mã sử dụng tìm kiếm tuần tự O(n) trên mảng động. '
        'Thêm sinh viên mới sử dụng thao tác push() với chi phí O(1) trung bình.', style='List Bullet'
    )
    doc.add_paragraph(
        'Read (Đọc): Hiển thị danh sách sinh viên dưới dạng bảng phân trang, xem chi tiết thông tin từng sinh viên. '
        'Hỗ trợ hiển thị tổng quan trên dashboard bao gồm: tổng số sinh viên, số sinh viên xuất sắc/giỏi/khá/yếu, '
        'và danh sách sinh viên mới thêm gần đây.', style='List Bullet'
    )
    doc.add_paragraph(
        'Update (Cập nhật): Cho phép sửa đổi thông tin sinh viên đã tồn tại. Hệ thống sẽ tìm sinh viên bằng mã '
        '(tìm kiếm tuần tự O(n)), cập nhật thông tin mới và validate lại trước khi lưu.', style='List Bullet'
    )
    doc.add_paragraph(
        'Delete (Xóa): Cho phép xóa sinh viên khỏi danh sách. Sử dụng thao tác removeAt() của mảng động, '
        'chi phí O(n) do cần dịch chuyển các phần tử sau vị trí xóa.', style='List Bullet'
    )
    
    doc.add_heading('2.2.2. Tìm kiếm', level=3)
    doc.add_paragraph(
        'Hệ thống hỗ trợ tìm kiếm sinh viên theo nhiều tiêu chí linh hoạt: tìm kiếm theo mã sinh viên, theo họ tên, '
        'theo lớp, hoặc tìm kiếm trên tất cả các trường. Ngoài ra, hệ thống còn hỗ trợ lọc theo xếp loại học lực '
        '(Xuất sắc, Giỏi, Khá, Trung bình, Yếu) và theo giới tính (Nam, Nữ). Thuật toán tìm kiếm chính sử dụng '
        'tìm kiếm tuần tự (Linear Search) kết hợp so sánh chuỗi con (string matching) thông qua hàm includes(). '
        'Bổ sung thuật toán tìm kiếm nhị phân (Binary Search) cho trường hợp mảng đã được sắp xếp.'
    )
    
    doc.add_heading('2.2.3. Sắp xếp', level=3)
    doc.add_paragraph(
        'Danh sách sinh viên có thể được sắp xếp theo nhiều tiêu chí: theo mã sinh viên, theo tên (lấy phần tên cuối cùng '
        'trong họ tên đầy đủ để sắp xếp theo đúng thông lệ tiếng Việt), theo điểm trung bình, hoặc theo tên lớp. '
        'Mỗi tiêu chí đều hỗ trợ sắp xếp tăng dần (ascending) và giảm dần (descending). Thuật toán sắp xếp chính là '
        'QuickSort với độ phức tạp trung bình O(n log n). Lưu ý quan trọng: mảng được nhân bản (clone) trước khi sắp xếp '
        'để không thay đổi thứ tự gốc của dữ liệu.'
    )
    
    doc.add_heading('2.2.4. Thống kê', level=3)
    doc.add_paragraph(
        'Hệ thống cung cấp chức năng thống kê tổng hợp với các chỉ số: tổng số sinh viên, điểm trung bình chung, '
        'điểm cao nhất, điểm thấp nhất, phân bố theo xếp loại (Xuất sắc/Giỏi/Khá/Trung bình/Yếu), phân bố theo giới tính, '
        'và thống kê chi tiết theo từng lớp (số lượng, điểm trung bình, phân bố xếp loại). Tất cả các chỉ số được tính toán '
        'chỉ trong MỘT LẦN duyệt mảng (thuật toán single-pass O(n)).'
    )
    
    # --- 2.3 ---
    doc.add_heading('2.3. Kiến trúc hệ thống', level=2)
    doc.add_paragraph(
        'Hệ thống được thiết kế theo kiến trúc module hóa, mỗi file JavaScript đảm nhận một nhiệm vụ riêng biệt, '
        'giúp mã nguồn dễ hiểu, dễ bảo trì và dễ mở rộng. Kiến trúc tổng thể bao gồm 9 file JavaScript chia thành '
        '3 nhóm chính:'
    )
    
    doc.add_paragraph('Nhóm 1 — Cấu trúc dữ liệu và thuật toán (4 file):')
    doc.add_paragraph('CauTrucMangDong.js: Lý thuyết và minh họa chi tiết về mảng động.', style='List Bullet')
    doc.add_paragraph('ThuatToanSapXep.js: Thuật toán QuickSort, BubbleSort và các hàm so sánh sinh viên.', style='List Bullet')
    doc.add_paragraph('ThuatToanTimKiem.js: Thuật toán tìm kiếm tuần tự, tìm kiếm nhị phân, tìm kiếm đa tiêu chí và lọc dữ liệu.', style='List Bullet')
    doc.add_paragraph('ThuatToanThongKe.js: Thuật toán tính GPA, xếp loại, thống kê tổng hợp single-pass và đếm phân nhóm.', style='List Bullet')
    
    doc.add_paragraph('Nhóm 2 — Mô hình dữ liệu và nghiệp vụ (3 file):')
    doc.add_paragraph('MangDong.js: Lớp DynamicArray — triển khai cấu trúc mảng động thực tế.', style='List Bullet')
    doc.add_paragraph('SinhVien.js: Lớp Student — mô hình hóa đối tượng sinh viên.', style='List Bullet')
    doc.add_paragraph('QuanLySinhVien.js: Lớp StudentManager — quản lý nghiệp vụ CRUD, tìm kiếm, sắp xếp, thống kê.', style='List Bullet')
    
    doc.add_paragraph('Nhóm 3 — Giao diện và điều khiển (2 file):')
    doc.add_paragraph('DieuKhienGiaoDien.js: Lớp UIController — điều khiển giao diện, xử lý sự kiện DOM.', style='List Bullet')
    doc.add_paragraph('UngDung.js: Lớp Application — điểm khởi đầu, kết nối và khởi tạo hệ thống.', style='List Bullet')
    
    # --- 2.4 ---
    doc.add_heading('2.4. Mô hình đối tượng', level=2)
    doc.add_paragraph(
        'Hệ thống sử dụng mô hình đối tượng gồm 5 lớp chính với mối quan hệ rõ ràng:'
    )
    
    # Bảng mô hình đối tượng
    table = doc.add_table(rows=6, cols=4)
    add_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Lớp', 'File', 'Thuộc tính chính', 'Phương thức chính']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')
    
    obj_data = [
        ['Student', 'SinhVien.js', 'id, name, dob, gender, className, email, phone, scores', 'calculateGPA(), getRank(), validate(), update()'],
        ['DynamicArray', 'MangDong.js', 'data, length, capacity', 'push(), pop(), get(), set(), insertAt(), removeAt(), sort(), filter(), clone()'],
        ['StudentManager', 'QuanLySinhVien.js', 'students (DynamicArray), storageKey', 'addStudent(), deleteStudent(), searchStudents(), sortStudents(), getStatistics()'],
        ['UIController', 'DieuKhienGiaoDien.js', 'manager, currentPage, itemsPerPage', 'init(), renderStudentTable(), handleSearch(), handleSort(), updateDashboard()'],
        ['Application', 'UngDung.js', 'manager, ui', 'start(), loadSampleData(), exportReport()'],
    ]
    
    for i, row_data in enumerate(obj_data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = cell_text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'Mối quan hệ giữa các lớp: Application chứa StudentManager và UIController. StudentManager sử dụng DynamicArray '
        'để lưu trữ danh sách các đối tượng Student. UIController tham chiếu đến StudentManager để truy cập dữ liệu và '
        'gọi các thuật toán. Các module thuật toán (ThuatToanSapXep, ThuatToanTimKiem, ThuatToanThongKe) được sử dụng '
        'bởi cả DynamicArray và StudentManager.'
    )
    
    # --- 2.5 ---
    doc.add_heading('2.5. Thiết kế cấu trúc dữ liệu', level=2)
    doc.add_paragraph(
        'Cấu trúc dữ liệu trung tâm của hệ thống là lớp DynamicArray, được triển khai tự viết (custom implementation) '
        'thay vì sử dụng mảng JavaScript tích hợp. Lý do chọn mảng động:'
    )
    doc.add_paragraph('Truy cập nhanh O(1): Phù hợp cho việc hiển thị danh sách, phân trang, và truy cập sinh viên theo vị trí.', style='List Bullet')
    doc.add_paragraph('Thêm phần tử hiệu quả O(1) trung bình: Thao tác thêm sinh viên mới rất nhanh nhờ cơ chế khấu hao.', style='List Bullet')
    doc.add_paragraph('Hỗ trợ sắp xếp O(n log n): Mảng liên tiếp trong bộ nhớ phù hợp với QuickSort (cache-friendly).', style='List Bullet')
    doc.add_paragraph('Dễ tuần tự hóa: Chuyển đổi sang JSON để lưu localStorage đơn giản hơn danh sách liên kết.', style='List Bullet')
    
    doc.add_paragraph(
        'Dung lượng ban đầu của mảng động được thiết lập là 20, phù hợp cho một lớp nhỏ. Khi số sinh viên vượt quá 20, '
        'mảng tự động mở rộng lên 40, rồi 80, 160... đảm bảo luôn đủ chỗ cho dữ liệu mới. Dữ liệu được lưu trữ '
        'bền vững (persistent) thông qua localStorage của trình duyệt, đảm bảo không mất dữ liệu khi đóng trang web.'
    )
    
    # --- 2.6 ---
    doc.add_heading('2.6. Luồng xử lý chính', level=2)
    
    doc.add_paragraph(
        'Luồng thêm sinh viên mới: Người dùng nhập thông tin vào form → UIController thu thập dữ liệu → '
        'Tạo đối tượng Student → Gọi StudentManager.addStudent() → Validate dữ liệu (Student.validate()) → '
        'Kiểm tra trùng mã (tìm kiếm tuần tự O(n)) → Thêm vào DynamicArray (push O(1)) → Lưu localStorage → '
        'Cập nhật giao diện.'
    )
    
    doc.add_paragraph(
        'Luồng tìm kiếm: Người dùng nhập từ khóa và chọn tiêu chí → UIController gọi StudentManager.searchStudents() → '
        'Sử dụng filter() của DynamicArray → Duyệt tuần tự O(n), kiểm tra từ khóa trên các trường (includes) → '
        'Lọc thêm theo xếp loại và giới tính → Trả về kết quả → Hiển thị lên bảng.'
    )
    
    doc.add_paragraph(
        'Luồng sắp xếp: Người dùng chọn tiêu chí sắp xếp → UIController gọi StudentManager.sortStudents() → '
        'Clone mảng O(n) → Tạo hàm so sánh (ThuatToanSapXep.taoHamSoSanh()) → Gọi DynamicArray.sort() → '
        'QuickSort O(n log n) → Trả về mảng đã sắp xếp → Hiển thị lên bảng phân trang.'
    )
    
    doc.add_paragraph(
        'Luồng thống kê: UIController gọi StudentManager.getStatistics() → Duyệt DynamicArray MỘT LẦN (single-pass O(n)) → '
        'Cập nhật đồng thời: tổng điểm, max, min, đếm xếp loại, đếm giới tính, đếm theo lớp → '
        'Tính trung bình → Trả về kết quả → Vẽ biểu đồ Chart.js.'
    )
    
    doc.add_page_break()
    
    # ================================================================
    #                    CHƯƠNG 3: THUẬT TOÁN SỬ DỤNG
    # ================================================================
    doc.add_heading('CHƯƠNG 3: THUẬT TOÁN SỬ DỤNG', level=1)
    
    doc.add_paragraph(
        'Chương này trình bày chi tiết các thuật toán được sử dụng trong hệ thống quản lý sinh viên. '
        'Mỗi thuật toán được phân tích kỹ lưỡng về nguyên lý hoạt động, mã giả (pseudocode), '
        'ví dụ minh họa cụ thể, và đánh giá độ phức tạp.'
    )
    
    # --- 3.1 ---
    doc.add_heading('3.1. Thuật toán trên mảng động', level=2)
    
    doc.add_heading('3.1.1. Thuật toán mở rộng mảng (_resize)', level=3)
    doc.add_paragraph(
        'Đây là thuật toán quan trọng nhất và đặc trưng nhất của mảng động. Khi số phần tử hiện tại (length) '
        'bằng dung lượng tối đa (capacity), mảng cần được mở rộng để chứa thêm phần tử mới.'
    )
    
    doc.add_paragraph('Mã giả:')
    p = doc.add_paragraph()
    run = p.add_run(
        'FUNCTION _resize():\n'
        '    dungLuongMoi = capacity × 2\n'
        '    mangMoi = TẠO MẢNG MỚI kích thước dungLuongMoi\n'
        '    FOR i = 0 TO length - 1:\n'
        '        mangMoi[i] = data[i]    // Copy từng phần tử\n'
        '    data = mangMoi\n'
        '    capacity = dungLuongMoi'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Ví dụ minh họa:')
    run.bold = True
    
    doc.add_paragraph(
        'Giả sử mảng đang chứa 4 sinh viên {SV001, SV002, SV003, SV004} với capacity = 4. '
        'Khi thêm sinh viên thứ 5 (SV005), hàm push() phát hiện length (4) ≥ capacity (4) → gọi _resize():'
    )
    
    p = doc.add_paragraph()
    run = p.add_run(
        'Trước resize:\n'
        '  data = [SV001][SV002][SV003][SV004]  (length=4, capacity=4)\n\n'
        '  Bước 1: dungLuongMoi = 4 × 2 = 8\n'
        '  Bước 2: Tạo mangMoi có 8 ô\n'
        '  Bước 3: Copy 4 phần tử sang mangMoi\n\n'
        'Sau resize:\n'
        '  data = [SV001][SV002][SV003][SV004][_][_][_][_]  (length=4, capacity=8)\n\n'
        'Bây giờ thêm SV005:\n'
        '  data = [SV001][SV002][SV003][SV004][SV005][_][_][_]  (length=5, capacity=8)'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph('Độ phức tạp: O(n) — cần copy n phần tử sang mảng mới.')
    
    doc.add_heading('3.1.2. Thuật toán thêm phần tử vào cuối (push)', level=3)
    doc.add_paragraph(
        'Thao tác push là thao tác được sử dụng thường xuyên nhất trong hệ thống — mỗi lần thêm sinh viên mới '
        'đều gọi push. Thuật toán kiểm tra xem mảng có đầy không, nếu đầy thì mở rộng, sau đó gán phần tử '
        'vào vị trí length và tăng length lên 1.'
    )
    
    doc.add_paragraph('Mã giả:')
    p = doc.add_paragraph()
    run = p.add_run(
        'FUNCTION push(phanTu):\n'
        '    IF length >= capacity:\n'
        '        _resize()              // Mở rộng nếu đầy\n'
        '    data[length] = phanTu      // Gán vào vị trí tiếp theo\n'
        '    length = length + 1        // Tăng số phần tử\n'
        '    RETURN length'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph('Độ phức tạp: O(1) trung bình (phân tích khấu hao), O(n) trường hợp xấu nhất (khi resize).')
    
    doc.add_heading('3.1.3. Thuật toán xóa phần tử tại vị trí (removeAt)', level=3)
    doc.add_paragraph(
        'Khi xóa sinh viên khỏi danh sách, thuật toán removeAt được sử dụng. Thao tác này cần dịch chuyển '
        'tất cả các phần tử sau vị trí xóa lên trước 1 ô để lấp đầy khoảng trống.'
    )
    
    doc.add_paragraph('Mã giả:')
    p = doc.add_paragraph()
    run = p.add_run(
        'FUNCTION removeAt(viTri):\n'
        '    IF viTri < 0 OR viTri >= length:\n'
        '        RETURN undefined       // Vị trí không hợp lệ\n'
        '    phanTuDaXoa = data[viTri]  // Lưu phần tử cần xóa\n'
        '    FOR i FROM viTri TO length - 2:\n'
        '        data[i] = data[i + 1]  // Dịch phần tử lên trước\n'
        '    data[length - 1] = undefined\n'
        '    length = length - 1\n'
        '    RETURN phanTuDaXoa'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Ví dụ minh họa:')
    run.bold = True
    
    doc.add_paragraph(
        'Xóa sinh viên tại vị trí index = 1 (SV002) trong mảng [SV001, SV002, SV003, SV004]:'
    )
    
    p = doc.add_paragraph()
    run = p.add_run(
        'Trước: [SV001][SV002][SV003][SV004]  (length=4)\n'
        '                 ↑ xóa vị trí 1\n\n'
        'Dịch:  data[1] = data[2]  →  [SV001][SV003][SV003][SV004]\n'
        '       data[2] = data[3]  →  [SV001][SV003][SV004][SV004]\n'
        '       data[3] = undefined\n\n'
        'Sau:   [SV001][SV003][SV004][_]  (length=3)'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph('Độ phức tạp: O(n) — xóa ở đầu phải dịch n-1 phần tử, xóa ở cuối là O(1).')
    
    doc.add_heading('3.1.4. Thuật toán chèn phần tử tại vị trí (insertAt)', level=3)
    doc.add_paragraph(
        'Thuật toán chèn phần tử tại một vị trí bất kỳ trong mảng. Các phần tử từ vị trí chèn đến cuối '
        'phải được dịch ra sau 1 ô để tạo chỗ trống. Lưu ý quan trọng: phải dịch TỪ CUỐI về đầu để '
        'tránh ghi đè dữ liệu.'
    )
    
    doc.add_paragraph('Mã giả:')
    p = doc.add_paragraph()
    run = p.add_run(
        'FUNCTION insertAt(viTri, phanTu):\n'
        '    IF viTri < 0 OR viTri > length: RETURN false\n'
        '    IF length >= capacity: _resize()\n'
        '    FOR i FROM length DOWNTO viTri + 1:\n'
        '        data[i] = data[i-1]    // Dịch ra sau (từ cuối!)\n'
        '    data[viTri] = phanTu        // Đặt phần tử mới\n'
        '    length = length + 1\n'
        '    RETURN true'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Ví dụ minh họa:')
    run.bold = True
    doc.add_paragraph(
        'Chèn sinh viên "SV_NEW" tại vị trí 1 trong mảng [SV001, SV003, SV004]:'
    )
    p = doc.add_paragraph()
    run = p.add_run(
        'Trước: [SV001][SV003][SV004][_]  (length=3, capacity=4)\n\n'
        'Dịch:  data[3] = data[2]  →  [SV001][SV003][SV004][SV004]\n'
        '       data[2] = data[1]  →  [SV001][SV003][SV003][SV004]\n\n'
        'Chèn:  data[1] = SV_NEW   →  [SV001][SV_NEW][SV003][SV004]\n\n'
        'Sau:   [SV001][SV_NEW][SV003][SV004]  (length=4)'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph('Độ phức tạp: O(n) — phải dịch các phần tử.')
    
    doc.add_heading('3.1.5. Thuật toán nhân bản mảng (clone)', level=3)
    doc.add_paragraph(
        'Thuật toán nhân bản tạo một bản sao hoàn toàn mới của mảng, giúp sắp xếp hiển thị mà không ảnh hưởng '
        'đến thứ tự gốc. Trong hệ thống, mỗi khi sắp xếp danh sách sinh viên, mảng gốc được clone trước, '
        'sau đó mới sắp xếp bản sao.'
    )
    
    doc.add_paragraph('Mã giả:')
    p = doc.add_paragraph()
    run = p.add_run(
        'FUNCTION clone():\n'
        '    banSao = TẠO DynamicArray MỚI (capacity = this.capacity)\n'
        '    FOR i = 0 TO length - 1:\n'
        '        banSao.push(data[i])   // Copy từng phần tử\n'
        '    RETURN banSao'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph('Độ phức tạp: O(n) — copy n phần tử.')
    
    # --- 3.2 ---
    doc.add_heading('3.2. Thuật toán sắp xếp QuickSort', level=2)
    
    doc.add_paragraph(
        'QuickSort là thuật toán sắp xếp chính được sử dụng trong hệ thống để sắp xếp danh sách sinh viên '
        'theo các tiêu chí: mã sinh viên, tên, điểm trung bình hoặc lớp. Thuật toán hoạt động theo chiến lược '
        '"chia để trị" (Divide and Conquer) với ba bước chính: chọn chốt (pivot), phân hoạch (partition), '
        'và đệ quy sắp xếp hai phần.'
    )
    
    doc.add_heading('3.2.1. Hàm đệ quy chính', level=3)
    doc.add_paragraph('Mã giả:')
    p = doc.add_paragraph()
    run = p.add_run(
        'FUNCTION quickSort(mang, viTriDau, viTriCuoi, hamSoSanh):\n'
        '    IF viTriDau >= viTriCuoi:\n'
        '        RETURN                  // Điều kiện dừng\n'
        '    viTriChot = phanHoach(mang, viTriDau, viTriCuoi, hamSoSanh)\n'
        '    quickSort(mang, viTriDau, viTriChot - 1, hamSoSanh)     // Nửa trái\n'
        '    quickSort(mang, viTriChot + 1, viTriCuoi, hamSoSanh)    // Nửa phải'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_heading('3.2.2. Hàm phân hoạch (Partition)', level=3)
    doc.add_paragraph(
        'Hàm phân hoạch là trái tim của QuickSort. Nó chọn phần tử cuối cùng làm chốt (pivot), sau đó duyệt '
        'mảng từ trái sang phải, đưa các phần tử nhỏ hơn hoặc bằng pivot về bên trái. Biến i đóng vai trò '
        'là "biên" của vùng chứa các phần tử ≤ pivot.'
    )
    
    doc.add_paragraph('Mã giả:')
    p = doc.add_paragraph()
    run = p.add_run(
        'FUNCTION phanHoach(mang, viTriDau, viTriCuoi, hamSoSanh):\n'
        '    chot = mang[viTriCuoi]      // Chọn phần tử cuối làm chốt\n'
        '    i = viTriDau - 1             // Biên vùng ≤ chốt\n'
        '    FOR j FROM viTriDau TO viTriCuoi - 1:\n'
        '        IF hamSoSanh(mang[j], chot) <= 0:\n'
        '            i = i + 1\n'
        '            SWAP(mang[i], mang[j])   // Đưa về vùng trái\n'
        '    SWAP(mang[i+1], mang[viTriCuoi]) // Đặt chốt vào đúng vị trí\n'
        '    RETURN i + 1                      // Trả về vị trí chốt'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Ví dụ minh họa chi tiết:')
    run.bold = True
    
    doc.add_paragraph(
        'Sắp xếp mảng sinh viên theo điểm trung bình tăng dần: [7.5, 9.0, 5.0, 8.0, 6.5]'
    )
    
    p = doc.add_paragraph()
    run = p.add_run(
        'Mảng ban đầu: [7.5, 9.0, 5.0, 8.0, 6.5]\n'
        'Chọn chốt = 6.5 (phần tử cuối)\n\n'
        '--- Phân hoạch ---\n'
        'i = -1, j = 0: mang[0]=7.5 > 6.5 → không đổi\n'
        'i = -1, j = 1: mang[1]=9.0 > 6.5 → không đổi\n'
        'i = -1, j = 2: mang[2]=5.0 ≤ 6.5 → i=0, swap(mang[0],mang[2])\n'
        '  → [5.0, 9.0, 7.5, 8.0, 6.5]\n'
        'i =  0, j = 3: mang[3]=8.0 > 6.5 → không đổi\n'
        'Đặt chốt: swap(mang[1], mang[4]) → [5.0, 6.5, 7.5, 8.0, 9.0]\n'
        'Vị trí chốt = 1\n\n'
        '--- Đệ quy nửa trái [5.0] → đã sắp xếp ---\n'
        '--- Đệ quy nửa phải [7.5, 8.0, 9.0] → đã sắp xếp ---\n\n'
        'Kết quả: [5.0, 6.5, 7.5, 8.0, 9.0] ✓'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_heading('3.2.3. Hàm so sánh sinh viên', level=3)
    doc.add_paragraph(
        'Hệ thống triển khai hàm tạo bộ so sánh (comparator factory) cho phép tạo hàm so sánh động '
        'theo nhiều tiêu chí khác nhau. Hàm taoHamSoSanh(truong, thuTu) nhận hai tham số: trường sắp xếp '
        '(id, name, gpa, class) và thứ tự (asc/desc), trả về hàm so sánh phù hợp.'
    )
    
    doc.add_paragraph('Đặc biệt với tiêu chí sắp xếp theo tên (name), hệ thống sử dụng kỹ thuật tách phần tên cuối '
        'cùng trong họ tên đầy đủ. Ví dụ: "Nguyễn Văn An" → lấy "An" để so sánh, phù hợp với quy ước sắp xếp '
        'tên tiếng Việt. Hàm localeCompare() được sử dụng với locale "vi" để so sánh chính xác ký tự tiếng Việt.')
    
    # --- 3.3 ---
    doc.add_heading('3.3. Thuật toán sắp xếp BubbleSort', level=2)
    doc.add_paragraph(
        'BubbleSort được triển khai trong hệ thống với mục đích minh họa và so sánh hiệu suất với QuickSort. '
        'Thuật toán duyệt mảng nhiều lượt, mỗi lượt so sánh hai phần tử liền kề và hoán đổi nếu sai thứ tự.'
    )
    
    doc.add_paragraph('Mã giả (có tối ưu cờ dừng sớm):')
    p = doc.add_paragraph()
    run = p.add_run(
        'FUNCTION bubbleSort(mang, hamSoSanh):\n'
        '    n = mang.length\n'
        '    FOR i FROM 0 TO n - 2:\n'
        '        daDoi = false\n'
        '        FOR j FROM 0 TO n - i - 2:\n'
        '            IF hamSoSanh(mang[j], mang[j+1]) > 0:\n'
        '                SWAP(mang[j], mang[j+1])\n'
        '                daDoi = true\n'
        '        IF daDoi == false:   // Không có hoán đổi → đã sắp xếp\n'
        '            BREAK             // Dừng sớm!'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Ví dụ minh họa:')
    run.bold = True
    
    doc.add_paragraph('Sắp xếp điểm TB: [7.5, 5.0, 9.0, 6.5]')
    
    p = doc.add_paragraph()
    run = p.add_run(
        'Lượt 1:\n'
        '  So sánh 7.5 và 5.0 → swap → [5.0, 7.5, 9.0, 6.5]\n'
        '  So sánh 7.5 và 9.0 → OK\n'
        '  So sánh 9.0 và 6.5 → swap → [5.0, 7.5, 6.5, 9.0]\n'
        '  → 9.0 "nổi" lên cuối\n\n'
        'Lượt 2:\n'
        '  So sánh 5.0 và 7.5 → OK\n'
        '  So sánh 7.5 và 6.5 → swap → [5.0, 6.5, 7.5, 9.0]\n'
        '  → 7.5 đúng vị trí\n\n'
        'Lượt 3:\n'
        '  So sánh 5.0 và 6.5 → OK\n'
        '  → Không có hoán đổi → DỪNG SỚM!\n\n'
        'Kết quả: [5.0, 6.5, 7.5, 9.0] ✓'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Nhận xét: BubbleSort có O(n²) trung bình, chậm hơn QuickSort O(n log n) đáng kể khi n lớn. '
        'Với n = 1000 sinh viên: BubbleSort cần khoảng 1.000.000 phép so sánh, QuickSort chỉ cần khoảng 10.000 phép. '
        'Tuy nhiên, BubbleSort có ưu điểm đơn giản và ổn định (stable sort).'
    )
    
    # --- 3.4 ---
    doc.add_heading('3.4. Thuật toán tìm kiếm tuần tự (Linear Search)', level=2)
    
    doc.add_paragraph(
        'Tìm kiếm tuần tự là thuật toán tìm kiếm mặc định trong hệ thống, được sử dụng cho nhiều chức năng: '
        'kiểm tra trùng mã sinh viên khi thêm mới, tìm sinh viên theo mã để cập nhật/xóa, và tìm kiếm '
        'theo từ khóa trên nhiều trường.'
    )
    
    doc.add_heading('3.4.1. Tìm kiếm đơn giản', level=3)
    doc.add_paragraph('Mã giả:')
    p = doc.add_paragraph()
    run = p.add_run(
        'FUNCTION timTuanTu(mang, dieuKien):\n'
        '    FOR i FROM 0 TO mang.length - 1:\n'
        '        IF dieuKien(mang[i]) == true:\n'
        '            RETURN i          // Tìm thấy → trả về vị trí\n'
        '    RETURN -1                  // Không tìm thấy'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Ví dụ minh họa:')
    run.bold = True
    doc.add_paragraph(
        'Tìm sinh viên có mã "SV003" trong mảng [SV001, SV002, SV003, SV004, SV005]:'
    )
    
    p = doc.add_paragraph()
    run = p.add_run(
        'Bước 1: i=0, kiểm tra SV001.id == "SV003"? → KHÔNG\n'
        'Bước 2: i=1, kiểm tra SV002.id == "SV003"? → KHÔNG\n'
        'Bước 3: i=2, kiểm tra SV003.id == "SV003"? → CÓ!\n'
        '→ Trả về vị trí 2\n'
        '→ Số bước: 3 (trong tổng 5 phần tử)'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_heading('3.4.2. Tìm kiếm đa tiêu chí', level=3)
    doc.add_paragraph(
        'Hệ thống mở rộng tìm kiếm tuần tự thành tìm kiếm đa tiêu chí: với mỗi sinh viên, kiểm tra xem '
        'từ khóa có xuất hiện trong mã sinh viên, họ tên, hoặc tên lớp. Thuật toán chuẩn hóa từ khóa (chuyển '
        'về chữ thường, xóa khoảng trắng thừa) rồi sử dụng hàm includes() để kiểm tra chuỗi con.'
    )
    
    doc.add_paragraph('Mã giả:')
    p = doc.add_paragraph()
    run = p.add_run(
        'FUNCTION timKiemSinhVien(danhSach, tuKhoa):\n'
        '    tuKhoaChuanHoa = tuKhoa.toLowerCase().trim()\n'
        '    IF tuKhoaChuanHoa rỗng: RETURN danhSach\n'
        '    ketQua = []\n'
        '    FOR mỗi sv TRONG danhSach:\n'
        '        IF sv.id    CHỨA tuKhoaChuanHoa\n'
        '        OR sv.name  CHỨA tuKhoaChuanHoa\n'
        '        OR sv.class CHỨA tuKhoaChuanHoa\n'
        '        OR sv.rank  CHỨA tuKhoaChuanHoa:\n'
        '            ketQua.THÊM(sv)\n'
        '    RETURN ketQua'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Ví dụ minh họa:')
    run.bold = True
    doc.add_paragraph(
        'Từ khóa: "nguyễn" → Chuẩn hóa: "nguyễn"\n'
        'Duyệt danh sách 30 sinh viên:\n'
        '- SV001 "Nguyễn Văn An" → name chứa "nguyễn" → ĐẠT\n'
        '- SV002 "Trần Thị Bình" → không trùng → BỎ QUA\n'
        '- ...\n'
        '- SV014 "Nguyễn Thị Hồng" → name chứa "nguyễn" → ĐẠT\n'
        '- SV026 "Nguyễn Thị Thảo" → name chứa "nguyễn" → ĐẠT\n'
        'Kết quả: 3 sinh viên phù hợp'
    )
    
    # --- 3.5 ---
    doc.add_heading('3.5. Thuật toán tìm kiếm nhị phân (Binary Search)', level=2)
    
    doc.add_paragraph(
        'Tìm kiếm nhị phân được triển khai trong hệ thống như một phương pháp tìm kiếm hiệu quả cao '
        'cho trường hợp mảng đã được sắp xếp. Thuật toán hoạt động bằng cách liên tục chia đôi không gian '
        'tìm kiếm, mỗi bước loại bỏ một nửa số phần tử không cần xét.'
    )
    
    doc.add_paragraph('Mã giả:')
    p = doc.add_paragraph()
    run = p.add_run(
        'FUNCTION timNhiPhan(mang, giaTri, hamSoSanh):\n'
        '    trai = 0\n'
        '    phai = mang.length - 1\n'
        '    WHILE trai <= phai:\n'
        '        giua = trai + (phai - trai) / 2    // Tránh tràn số\n'
        '        ketQua = hamSoSanh(mang[giua], giaTri)\n'
        '        IF ketQua == 0:\n'
        '            RETURN giua                      // Tìm thấy!\n'
        '        ELSE IF ketQua < 0:\n'
        '            trai = giua + 1                  // Tìm nửa PHẢI\n'
        '        ELSE:\n'
        '            phai = giua - 1                  // Tìm nửa TRÁI\n'
        '    RETURN -1                                // Không tìm thấy'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Ví dụ minh họa chi tiết:')
    run.bold = True
    
    doc.add_paragraph(
        'Tìm sinh viên có điểm TB = 8.0 trong mảng ĐÃ SẮP XẾP theo điểm TB tăng dần:'
    )
    
    p = doc.add_paragraph()
    run = p.add_run(
        'Mảng: [4.5, 5.0, 6.0, 6.5, 7.0, 7.5, 8.0, 8.5, 9.0, 9.5]\n'
        'Chỉ số: [0]  [1]  [2]  [3]  [4]  [5]  [6]  [7]  [8]  [9]\n'
        'Tìm giá trị: 8.0\n\n'
        'Bước 1: trai=0, phai=9, giua=(0+9)/2=4\n'
        '  mang[4] = 7.0 < 8.0 → tìm nửa phải: trai = 5\n\n'
        'Bước 2: trai=5, phai=9, giua=(5+9)/2=7\n'
        '  mang[7] = 8.5 > 8.0 → tìm nửa trái: phai = 6\n\n'
        'Bước 3: trai=5, phai=6, giua=(5+6)/2=5\n'
        '  mang[5] = 7.5 < 8.0 → tìm nửa phải: trai = 6\n\n'
        'Bước 4: trai=6, phai=6, giua=(6+6)/2=6\n'
        '  mang[6] = 8.0 == 8.0 → TÌM THẤY tại vị trí 6!\n\n'
        '→ Chỉ cần 4 bước (thay vì 7 bước nếu dùng tìm kiếm tuần tự)\n'
        '→ Với 10 phần tử: log₂(10) ≈ 3.3, tối đa 4 bước ✓'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Lưu ý: Công thức tính chỉ số giữa sử dụng giua = trai + (phai - trai) / 2 thay vì giua = (trai + phai) / 2 '
        'để tránh tràn số khi trai + phai vượt quá giá trị tối đa của kiểu số. Đây là kỹ thuật lập trình '
        'an toàn được khuyến nghị.'
    )
    
    # --- 3.6 ---
    doc.add_heading('3.6. Thuật toán thống kê Single-pass', level=2)
    
    doc.add_paragraph(
        'Thuật toán thống kê tổng hợp của hệ thống là một minh chứng xuất sắc cho kỹ thuật "duyệt một lần" '
        '(single-pass). Thay vì duyệt mảng nhiều lần để thu thập từng chỉ số riêng, thuật toán chỉ duyệt MỘT LẦN '
        'DUY NHẤT nhưng cập nhật đồng thời tất cả các biến thống kê.'
    )
    
    doc.add_paragraph('Mã giả:')
    p = doc.add_paragraph()
    run = p.add_run(
        'FUNCTION thongKeTongHop(danhSach):\n'
        '    // Khởi tạo\n'
        '    tongDiem = 0\n'
        '    diemMax = -∞, diemMin = +∞\n'
        '    phanLoai = {xuatSac:0, gioi:0, kha:0, trungBinh:0, yeu:0}\n'
        '    theoLop = {}, theoGioiTinh = {}\n\n'
        '    // DUYỆT MỘT LẦN DUY NHẤT\n'
        '    FOR mỗi sv TRONG danhSach:\n'
        '        diemTB = sv.tinhDiemTB()\n'
        '        tongDiem += diemTB\n'
        '        IF diemTB > diemMax: diemMax = diemTB\n'
        '        IF diemTB < diemMin: diemMin = diemTB\n'
        '        // Phân loại học lực\n'
        '        IF diemTB >= 9.0: phanLoai.xuatSac++\n'
        '        ELSE IF diemTB >= 8.0: phanLoai.gioi++\n'
        '        ELSE IF diemTB >= 6.5: phanLoai.kha++\n'
        '        ELSE IF diemTB >= 5.0: phanLoai.trungBinh++\n'
        '        ELSE: phanLoai.yeu++\n'
        '        // Đếm theo lớp\n'
        '        theoLop[sv.lop]++\n'
        '        // Đếm theo giới tính\n'
        '        theoGioiTinh[sv.gioiTinh]++\n\n'
        '    diemTBChung = tongDiem / danhSach.length\n'
        '    RETURN {diemTBChung, diemMax, diemMin, phanLoai, theoLop, theoGioiTinh}'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Ví dụ minh họa:')
    run.bold = True
    
    doc.add_paragraph(
        'Cho danh sách 5 sinh viên với điểm TB lần lượt: 9.2, 8.0, 6.5, 4.5, 7.5'
    )
    
    p = doc.add_paragraph()
    run = p.add_run(
        'Duyệt:\n'
        '  SV1 (9.2): tongDiem=9.2, max=9.2, min=9.2, xuatSac=1\n'
        '  SV2 (8.0): tongDiem=17.2, max=9.2, min=8.0, gioi=1\n'
        '  SV3 (6.5): tongDiem=23.7, max=9.2, min=6.5, kha=1\n'
        '  SV4 (4.5): tongDiem=28.2, max=9.2, min=4.5, yeu=1\n'
        '  SV5 (7.5): tongDiem=35.7, max=9.2, min=4.5, kha=2\n\n'
        'Kết quả:\n'
        '  Điểm TB chung = 35.7 / 5 = 7.14\n'
        '  Điểm cao nhất = 9.2\n'
        '  Điểm thấp nhất = 4.5\n'
        '  Phân loại: Xuất sắc=1, Giỏi=1, Khá=2, TB=0, Yếu=1\n\n'
        '→ Tất cả chỉ với 1 lần duyệt (5 bước) thay vì 5 lần (25 bước)!'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Ngoài ra, thuật toán đếm phân nhóm sử dụng cấu trúc bảng băm (hash map) thông qua object JavaScript. '
        'Với mỗi sinh viên, lấy tên lớp (hoặc xếp loại, giới tính) làm khóa, tăng bộ đếm tương ứng lên 1. '
        'Chi phí: O(n) thời gian, O(k) bộ nhớ phụ (k = số nhóm, rất nhỏ).'
    )
    
    # --- 3.7 ---
    doc.add_heading('3.7. Thuật toán phân loại học lực', level=2)
    
    doc.add_paragraph(
        'Thuật toán phân loại học lực sử dụng cấu trúc rẽ nhánh if-else với các ngưỡng (threshold) '
        'được sắp xếp giảm dần. Việc sắp xếp giảm dần giúp tối ưu: sinh viên xuất sắc chỉ cần 1 phép so sánh, '
        'sinh viên yếu cần 4 phép so sánh, nhưng trung bình thì chỉ cần khoảng 2-3 phép so sánh.'
    )
    
    doc.add_paragraph('Mã giả:')
    p = doc.add_paragraph()
    run = p.add_run(
        'FUNCTION xepLoaiHocLuc(diemTB):\n'
        '    IF diemTB >= 9.0: RETURN "Xuất sắc"\n'
        '    IF diemTB >= 8.0: RETURN "Giỏi"\n'
        '    IF diemTB >= 6.5: RETURN "Khá"\n'
        '    IF diemTB >= 5.0: RETURN "Trung bình"\n'
        '    RETURN "Yếu"'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Ví dụ minh họa:')
    run.bold = True
    
    # Bảng ví dụ xếp loại
    table = doc.add_table(rows=6, cols=4)
    add_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Sinh viên', 'Điểm TB', 'Xếp loại', 'Số phép so sánh']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')
    
    rank_data = [
        ['Nguyễn Văn An', '9.17', 'Xuất sắc', '1'],
        ['Trần Thị Bình', '8.50', 'Giỏi', '2'],
        ['Lê Hoàng Cường', '7.50', 'Khá', '3'],
        ['Hoàng Văn Em', '6.00', 'Trung bình', '4'],
        ['Đỗ Minh Quang', '4.83', 'Yếu', '5 (tất cả)'],
    ]
    
    for i, row_data in enumerate(rank_data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = cell_text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
                    r.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph('Độ phức tạp: O(1) — tối đa 5 phép so sánh, không phụ thuộc vào kích thước dữ liệu.')
    
    # --- 3.8 ---
    doc.add_heading('3.8. Thuật toán Validation dữ liệu', level=2)
    
    doc.add_paragraph(
        'Thuật toán kiểm tra tính hợp lệ của dữ liệu sinh viên thực hiện kiểm tra tuần tự từng trường thông tin. '
        'Mỗi trường được áp dụng một quy tắc kiểm tra riêng, và tất cả lỗi được thu thập vào một mảng để '
        'hiển thị đồng thời cho người dùng (thay vì dừng ở lỗi đầu tiên).'
    )
    
    doc.add_paragraph('Các quy tắc kiểm tra:')
    doc.add_paragraph('Mã sinh viên: Không được để trống (kiểm tra chuỗi rỗng hoặc chỉ chứa khoảng trắng).', style='List Bullet')
    doc.add_paragraph('Họ và tên: Không được để trống.', style='List Bullet')
    doc.add_paragraph('Ngày sinh: Không được để trống và không được ở tương lai (so sánh với ngày hiện tại).', style='List Bullet')
    doc.add_paragraph('Giới tính: Phải là "Nam" hoặc "Nữ" (kiểm tra thuộc tập giá trị hợp lệ).', style='List Bullet')
    doc.add_paragraph('Lớp: Không được để trống.', style='List Bullet')
    doc.add_paragraph('Email: Nếu có, phải khớp biểu thức chính quy (regex): ^[^\\s@]+@[^\\s@]+\\.[^\\s@]+$', style='List Bullet')
    doc.add_paragraph('Điểm số (Toán, Văn, Anh): Nếu có, phải là số thực trong khoảng [0, 10].', style='List Bullet')
    
    doc.add_paragraph('Mã giả:')
    p = doc.add_paragraph()
    run = p.add_run(
        'FUNCTION validate():\n'
        '    loi = []                          // Mảng lỗi\n'
        '    IF id rỗng: loi.THÊM("Mã SV không được để trống")\n'
        '    IF name rỗng: loi.THÊM("Họ tên không được để trống")\n'
        '    IF dob rỗng: loi.THÊM("Ngày sinh không được để trống")\n'
        '    ELSE IF dob > ngayHienTai: loi.THÊM("Ngày sinh không hợp lệ")\n'
        '    IF gender KHÔNG THUỘC {"Nam","Nữ"}: loi.THÊM("Giới tính không hợp lệ")\n'
        '    IF className rỗng: loi.THÊM("Lớp không được để trống")\n'
        '    IF email CÓ GIÁ TRỊ AND email KHÔNG KHỚP regex:\n'
        '        loi.THÊM("Email không hợp lệ")\n'
        '    CHO MỖI môn TRONG {Toán, Văn, Anh}:\n'
        '        IF điểm CÓ GIÁ TRỊ AND (điểm < 0 OR điểm > 10):\n'
        '            loi.THÊM("Điểm {môn} phải từ 0 đến 10")\n'
        '    RETURN {valid: loi.length == 0, errors: loi}'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph('Độ phức tạp: O(1) — số trường kiểm tra cố định, không phụ thuộc vào kích thước dữ liệu.')
    
    # --- 3.9 ---
    doc.add_heading('3.9. Bảng tổng hợp độ phức tạp các thuật toán', level=2)
    
    doc.add_paragraph(
        'Bảng dưới đây tổng hợp độ phức tạp thời gian và bộ nhớ phụ của tất cả các thuật toán '
        'được sử dụng trong hệ thống quản lý sinh viên:'
    )
    
    # Bảng tổng hợp
    table = doc.add_table(rows=15, cols=5)
    add_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Thuật toán / Thao tác', 'Tốt nhất', 'Trung bình', 'Xấu nhất', 'Bộ nhớ phụ']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')
    
    complexity_data = [
        ['push (thêm cuối)', 'O(1)', 'O(1)*', 'O(n)', 'O(n)'],
        ['pop (xóa cuối)', 'O(1)', 'O(1)', 'O(1)', 'O(1)'],
        ['get / set (truy cập)', 'O(1)', 'O(1)', 'O(1)', 'O(1)'],
        ['insertAt (chèn)', 'O(1)', 'O(n)', 'O(n)', 'O(1)'],
        ['removeAt (xóa)', 'O(1)', 'O(n)', 'O(n)', 'O(1)'],
        ['clone (nhân bản)', 'O(n)', 'O(n)', 'O(n)', 'O(n)'],
        ['resize (mở rộng)', 'O(n)', 'O(n)', 'O(n)', 'O(n)'],
        ['QuickSort', 'O(n log n)', 'O(n log n)', 'O(n²)', 'O(log n)'],
        ['BubbleSort', 'O(n)', 'O(n²)', 'O(n²)', 'O(1)'],
        ['Tìm kiếm tuần tự', 'O(1)', 'O(n)', 'O(n)', 'O(1)'],
        ['Tìm kiếm nhị phân', 'O(1)', 'O(log n)', 'O(log n)', 'O(1)'],
        ['Thống kê single-pass', 'O(n)', 'O(n)', 'O(n)', 'O(k)'],
        ['Phân loại học lực', 'O(1)', 'O(1)', 'O(1)', 'O(1)'],
        ['Validation', 'O(1)', 'O(1)', 'O(1)', 'O(1)'],
    ]
    
    for i, row_data in enumerate(complexity_data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = cell_text
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph('Ghi chú: * = phân tích khấu hao (amortized analysis); k = số nhóm phân loại (rất nhỏ, thường < 10).')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Tổng kết chương 3: Hệ thống quản lý sinh viên sử dụng tổng cộng 8 nhóm thuật toán chính, '
        'bao phủ đầy đủ các thao tác cần thiết cho một ứng dụng quản lý dữ liệu. '
        'Các thuật toán được lựa chọn cân bằng giữa hiệu suất và độ đơn giản: QuickSort O(n log n) cho sắp xếp, '
        'tìm kiếm tuần tự O(n) cho dữ liệu không sắp xếp, tìm kiếm nhị phân O(log n) cho dữ liệu đã sắp xếp, '
        'và thống kê single-pass O(n) để thu thập đồng thời nhiều chỉ số. '
        'Cấu trúc mảng động với cơ chế resize gấp đôi đảm bảo thao tác push đạt O(1) trung bình, '
        'phù hợp cho việc thêm sinh viên mới liên tục.'
    )
    
    # ======================== KẾT LUẬN ========================
    doc.add_page_break()
    doc.add_heading('KẾT LUẬN', level=1)
    
    doc.add_paragraph(
        'Bài tập lớn đã hoàn thành việc xây dựng một hệ thống quản lý sinh viên hoàn chỉnh sử dụng '
        'cấu trúc dữ liệu mảng động và giao diện đồ họa web. Hệ thống được phát triển bằng JavaScript thuần '
        'với kiến trúc module hóa rõ ràng, bao gồm 9 file JavaScript phân chia theo chức năng.'
    )
    
    doc.add_paragraph(
        'Về mặt cấu trúc dữ liệu, hệ thống triển khai thành công lớp DynamicArray (Mảng động) tự viết '
        'với đầy đủ các thao tác: push, pop, get, set, insertAt, removeAt, filter, sort, binarySearch, clone. '
        'Cơ chế resize gấp đôi dung lượng đảm bảo thao tác thêm phần tử đạt O(1) trung bình theo phân tích khấu hao.'
    )
    
    doc.add_paragraph(
        'Về mặt thuật toán, hệ thống áp dụng đa dạng các thuật toán kinh điển: QuickSort O(n log n) cho sắp xếp, '
        'BubbleSort O(n²) cho minh họa, tìm kiếm tuần tự O(n) cho dữ liệu chưa sắp xếp, tìm kiếm nhị phân '
        'O(log n) cho dữ liệu đã sắp xếp, thống kê single-pass O(n) cho hiệu quả cao, và thuật toán phân loại '
        'dựa trên ngưỡng O(1) cho xếp loại học lực.'
    )
    
    doc.add_paragraph(
        'Về mặt chức năng, hệ thống đáp ứng đầy đủ các nghiệp vụ quản lý sinh viên: thêm mới với validate, '
        'cập nhật, xóa, hiển thị danh sách phân trang, tìm kiếm đa tiêu chí, sắp xếp đa trường, thống kê '
        'tổng hợp với biểu đồ trực quan, xuất/nhập dữ liệu JSON, và lưu trữ bền vững qua localStorage. '
        'Giao diện được thiết kế hiện đại, trực quan với HTML/CSS responsive.'
    )
    
    # ======================== LƯU FILE ========================
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'BaoCao_BTL_LapTrinhNangCao.docx')
    doc.save(output_path)
    print(f'Đã tạo báo cáo thành công: {output_path}')
    return output_path

if __name__ == '__main__':
    create_report()
